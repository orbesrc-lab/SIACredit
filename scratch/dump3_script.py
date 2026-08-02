import docx

filepath = 'c:/SIAC/f_5dc464440c50.docx'
outpath = 'c:/SIAC/scratch/dump3.txt'

try:
    doc = docx.Document(filepath)
    with open(outpath, 'w', encoding='utf-8') as f:
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                f.write(f"P{i}: {text}\n")
        
        f.write("\n--- TABLES ---\n")
        for t_idx, table in enumerate(doc.tables):
            f.write(f"\nTABLE {t_idx}:\n")
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                f.write(f"  R{r_idx}: " + " | ".join(cells) + "\n")
except Exception as e:
    print(f"Failed {filepath}: {e}")

print("Done dumping dump3.")
