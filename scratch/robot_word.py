import os
import docx
from supabase import create_client, Client
from dotenv import load_dotenv

load_dotenv()

url: str = os.getenv("SUPABASE_URL")
key: str = os.getenv("SUPABASE_KEY")
supabase: Client = create_client(url, key)

def extract_and_inject():
    filepath = 'c:/SIAC/es.docx'
    print(f"Abriendo documento {filepath}...")
    doc = docx.Document(filepath)
    
    print("Limpiando diccionario actual...")
    # Para evitar duplicados, limpiamos (opcional, pero util para no repetir)
    supabase.table('skel_diccionario_competencias').delete().neq('id', '00000000-0000-0000-0000-000000000000').execute()
    
    # 1. Extraer Competencias (Tablas que tienen "Competencia" en la celda 0,0)
    competencias_agregadas = {}
    
    for t_idx, table in enumerate(doc.tables):
        if not table.rows: continue
        header = table.rows[0].cells[0].text.strip().lower()
        if "competencia" in header:
            print(f"  -> Encontrada tabla de competencias (Índice {t_idx})")
            for row in table.rows[1:]:
                comp_name = row.cells[0].text.strip()
                if comp_name and comp_name not in competencias_agregadas:
                    res = supabase.table('skel_diccionario_competencias').insert({
                        "nombre": comp_name,
                        "tipo": "Evaluación"
                    }).execute()
                    if res.data:
                        competencias_agregadas[comp_name] = res.data[0]['id']
                        
    print(f"Se inyectaron {len(competencias_agregadas)} competencias exitosamente.")
    
    # 2. Extraer Comportamientos (Afirmaciones)
    # Algunas tablas tienen "Afirmación" o "Posible causa"
    
    # Vamos a crear una competencia "Generica" para colgarle las afirmaciones del diagnostico
    res_gen = supabase.table('skel_diccionario_competencias').insert({
        "nombre": "Diagnóstico General",
        "tipo": "Diagnóstico"
    }).execute()
    gen_id = res_gen.data[0]['id']
    
    comportamientos_count = 0
    for t_idx, table in enumerate(doc.tables):
        if not table.rows: continue
        header = table.rows[0].cells[1].text.strip().lower() if len(table.rows[0].cells) > 1 else ""
        if "afirmación" in header or "posible causa" in header:
            print(f"  -> Encontrada tabla de afirmaciones (Índice {t_idx})")
            for row in table.rows[1:]:
                # Las afirmaciones suelen estar en la celda 1
                if len(row.cells) > 1:
                    afirmacion = row.cells[1].text.strip()
                    if afirmacion:
                        supabase.table('skel_diccionario_comportamientos').insert({
                            "competencia_id": gen_id,
                            "descripcion": afirmacion,
                            "nivel_esperado": 5
                        }).execute()
                        comportamientos_count += 1
                        
    print(f"Se inyectaron {comportamientos_count} comportamientos/afirmaciones exitosamente.")
    print("¡Extracción e inyección finalizada!")

if __name__ == "__main__":
    extract_and_inject()
