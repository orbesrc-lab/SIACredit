import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_PATH = r"c:\SIAC\Manual_Usuario_SKEL.docx"

print("Iniciando la creacion del documento de Word para SKEL...")

doc = Document()

# Configurar Estilos Globales
style_normal = doc.styles['Normal']
font = style_normal.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(15, 23, 42) # slate-900

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(30, 58, 138) # Dark blue #1e3a8a
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(51, 65, 85) # slate-700
    return p

def add_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 116, 139) # slate-500
    return p

def add_paragraph(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
    p.add_run(text)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
    p.add_run(text)
    return p

def add_numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.bold = True
    p.add_run(text)
    return p

def add_callout(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("NOTA IMPORTANTE: " + text)
    run.italic = True
    run.font.color.rgb = RGBColor(180, 83, 9) # Amber-700
    return p

# --- CONSTRUIR DOCUMENTO ---

# Encabezado de Portada
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(40)
p_title.paragraph_format.space_after = Pt(10)
r_title = p_title.add_run("MANUAL DE USUARIO SKEL")
r_title.bold = True
r_title.font.size = Pt(26)
r_title.font.color.rgb = RGBColor(30, 58, 138)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(40)
r_sub = p_sub.add_run("Guía de Gestión de Calidad, Autoevaluación y Planes de Mejoramiento")
r_sub.italic = True
r_sub.font.size = Pt(14)
r_sub.font.color.rgb = RGBColor(71, 85, 105)

# Separador de sección
doc.add_page_break()

# 1. Conceptos Básicos
add_heading_1("1. Conceptos Básicos que Debes Conocer")
add_paragraph("Antes de empezar a interactuar con el sistema SKEL, es fundamental familiarizarse con la terminología básica del Modelo de Acreditación de Alta Calidad colombiano:")
add_bullet(" La universidad, escuela o corporación en la cual se aplica el sistema (ej. Corporación Universitaria Centro Superior - UNICUCES).", "Institución:")
add_bullet(" La carrera profesional o tecnológica que está realizando el proceso de autoevaluación (ej. Contaduría Pública).", "Programa Académico:")
add_bullet(" Las grandes dimensiones que estructuran la calidad académica e institucional (el modelo nacional cuenta con 12 Factores, tales como Profesores, Estudiantes, o Investigación).", "Factores:")
add_bullet(" Los aspectos particulares en los cuales se desglosa cada Factor. Son los ítems directos que reciben calificación de 1.0 a 5.0 (ej. Estatuto Profesoral, Misión).", "Características:")
add_bullet(" Archivos en formato PDF, Word o imágenes que actúan como soporte verificable de que los juicios de valor y calificaciones asignados son verídicos.", "Evidencias:")

# 2. Acceso
add_heading_1("2. Acceso e Inicio de Sesión")
add_numbered(" Ingrese al navegador de su preferencia (Chrome, Edge, Firefox) y navegue a la URL provista por su institución.")
add_numbered(" En el portal de acceso, introduzca su correo electrónico corporativo registrado y su contraseña.")
add_numbered(" Al iniciar sesión, el sistema adaptará su barra de herramientas y menús según uno de los tres perfiles disponibles: Superadministrador (gestión global), Administrador de Institución (gestión del programa y asignaciones) o Líder de Factor (calificación y evidencias).")

# 3. Configuración Inicial
add_heading_1("3. Paso 1: Configuración Inicial del Sistema (Rol: Administrador)")
add_paragraph("El Administrador de la Institución es el encargado de habilitar el entorno de trabajo para el programa académico.")
add_heading_2("A. Selección de Contexto Académico")
add_numbered(" Ingrese a la sección Configuración en el menú izquierdo.")
add_numbered(" En la parte superior de la página, elija la Institución y el Programa Académico (ej. Contaduría Pública).")
add_numbered(" Presione Guardar Contexto. Esto garantiza que todos los datos y reportes generados a partir de ahora queden clasificados en dicho programa de forma aislada.")
add_heading_2("B. Selección de Periodos y Pesos")
add_paragraph("En la misma pantalla, defina el año fiscal y lectivo correspondiente al proceso de autoevaluación actual (ej. 2026) y configure los pesos correspondientes a cada ponderación si el comité de currículo lo requiere.")
add_heading_2("C. Delegación de Responsables (Líderes de Factor)")
add_paragraph("Usted debe designar qué docente se encargará de evaluar cada Factor:")
add_numbered(" Diríjase a la sección Asignación de Líderes de Factor.")
add_numbered(" Ubique el Factor correspondiente (ej. Factor 3: Profesores) y escriba el correo electrónico del docente encargado.")
add_numbered(" Haga clic en Guardar Asignaciones. Automáticamente se habilitarán los accesos específicos para ese correo.")

# 4. Autoevaluación y Evidencias
add_heading_1("4. Paso 2: Autoevaluación e Integración del Módulo de Evidencias")
add_paragraph("La autoevaluación unifica el juicio cuantitativo y cualitativo de la institución con los soportes reales (evidencias) que lo sustentan.")
add_heading_2("A. Calificación de Características")
add_numbered(" El Líder de Factor ingresa al módulo Autoevaluación.")
add_numbered(" Despliega el factor asignado y selecciona la característica a calificar.")
add_numbered(" Selecciona una nota del 1.0 al 5.0 (escala Likert) y redacta la justificación cualitativa detallando fortalezas y debilidades.")
add_heading_2("B. Vinculación Automática con Evidencias")
add_callout("SKEL impide asignar calificaciones sobresalientes sin evidencias que las respalden.")
add_paragraph("Para subir y vincular evidencias, siga este orden:")
add_numbered(" Acceda al módulo Evidencias en la barra superior.")
add_numbered(" Seleccione el Factor y Característica correspondiente y arrastre los archivos PDF soporte. Dé clic en Subir.")
add_numbered(" Regrese al módulo Autoevaluación. Ingrese a la característica calificada: verá que el archivo PDF subido ahora aparece listado automáticamente al final del formulario de calificación, disponible para consulta de pares académicos y auditores.")

# 5. Encuestas
add_heading_1("5. Paso 3: Diseño y Aplicación de Encuestas a la Comunidad")
add_paragraph("Las encuestas capturan la percepción de la comunidad (estudiantes, egresados, profesores y administrativos).")
add_heading_2("A. Creación de la Encuesta")
add_numbered(" Ingrese al módulo Encuestas y presione Crear Nueva Encuesta.")
add_numbered(" Defina el título y la población objetivo (ej. Estudiantes).")
add_numbered(" Añada preguntas utilizando diversos tipos de respuesta: Sí/No, Selección Múltiple, Carga de Archivos de soporte y Calificación/Rating (Likert 1-5).")
add_numbered(" Presione Guardar Encuesta y actívela cambiándole el estado a Activo.")
add_heading_2("B. Recolección de Respuestas")
add_paragraph("Comparta el enlace generado por el sistema. Los resultados y sus diagramas de tortas y barras se actualizarán en tiempo real a medida que la comunidad responda.")

# 6. Planes de Mejoramiento
add_heading_1("6. Paso 4: Formulación y Desarrollo del Plan de Mejoramiento (Fase 3)")
add_paragraph("Cuando una característica obtiene una calificación baja (generalmente menor a 3.5), el sistema exige formular una Acción de Mejora.")
add_heading_2("A. Registro de la Acción de Mejora")
add_paragraph("En la parte inferior de la característica evaluada, haga clic en ➕ Nueva Acción y rellene las 9 variables de control:")
add_numbered(" Actividad: Qué acción concreta se va a realizar.", "1.")
add_numbered(" Meta: Qué resultado cuantificable se quiere lograr.", "2.")
add_numbered(" Fecha de Inicio: Cuándo arranca la actividad.", "3.")
add_numbered(" Fecha de Finalización: Fecha de vencimiento máxima.", "4.")
add_numbered(" Presupuesto en Tiempo: Horas o semanas estimadas.", "5.")
add_numbered(" Presupuesto Financiero: Inversión monetaria en pesos.", "6.")
add_numbered(" Rol del Responsable: LÍDER, ADMINISTRADOR u OPERATIVO.", "7.")
add_numbered(" Email del Responsable: Correo del encargado de ejecutar la acción.", "8.")
add_numbered(" Configuración del Indicador: Cómo se medirá el avance.", "9.")
add_heading_2("B. Explicación de los Tres Tipos de Indicadores")
add_heading_3("Tipo 1: Porcentaje (% de avance manual)")
add_paragraph("El avance se calcula dividiendo un Numerador (avance actual) entre un Denominador (meta total) multiplicado por 100.")
add_paragraph("Ejemplo: Si la meta es dictar 5 talleres de capacitación (denominador = 5) y se han dictado 3 (numerador = 3), el sistema calcula automáticamente un 60% de avance y cambia el estado a 'En proceso'.", "Práctica:")
add_heading_3("Tipo 2: Soporte Documental (0% o 100%)")
add_paragraph("El avance permanece en 0% (Pendiente) y pasa automáticamente a 100% (Completado) únicamente cuando el responsable sube el archivo PDF o imagen que certifique la culminación de la actividad en la plataforma.")
add_paragraph("Ejemplo: Subir la resolución firmada del nuevo Plan de Estudios. En el momento en que se carga, el plan se marca completado.", "Práctica:")
add_heading_3("Tipo 3: Opinión (Vínculo dinámico a Encuestas)")
add_paragraph("El avance se calcula dinámicamente según el promedio de respuestas que reciba una pregunta de escala 1 a 5 de una encuesta activa del sistema.")
add_paragraph("Ejemplo: Vincular la acción con la pregunta '¿Cómo califica el wifi?' en la encuesta de estudiantes. Si el promedio de calificación de las respuestas de los estudiantes es 4.5 sobre 5.0, el avance se proyecta automáticamente en un 90% en tiempo real.", "Práctica:")

# 7. Informes
add_heading_1("7. Paso 5: Generación y Análisis de Informes")
add_paragraph("Al finalizar la recolección, el sistema unifica los datos en un informe integral de calidad académica.")
add_heading_2("A. Visualización del Reporte")
add_paragraph("Al generar el informe, se visualizarán:")
add_bullet(" Gráfico interactivo que compara las notas de los factores con el estándar ideal de acreditación.")
add_bullet(" Muestra el presupuesto financiero acumulado de todas las acciones del programa.")
add_bullet(" Resume las horas de trabajo programadas para la mejora continua.")
add_bullet(" Listado consolidado interactivo con todos los planes de mejora.")

add_heading_2("B. Exportación y Descarga")
add_paragraph("Presione el botón Exportar Plan (Excel) para obtener la hoja de cálculo (.xls) premium con badges de estado estilizados. También puede presionar Ctrl + P para imprimir el informe en PDF con formato editorial optimizado.")

# Guardar el documento
doc.save(OUTPUT_PATH)
print(f"Documento de Word guardado con éxito en: {OUTPUT_PATH}")
