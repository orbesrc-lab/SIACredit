import docx

def dump_docx(filepath, outpath):
    try:
        doc = docx.Document(filepath)
        with open(outpath, 'w', encoding='utf-8') as f:
            for i, p in enumerate(doc.paragraphs):
                text = p.text.strip()
                if text:
                    f.write(f"P{i}: {text}\n")
            
            f.write("\n--- TABLES ---\n")
            for t_idx, table in enumerate(doc.tables):
                f.write(f"\nTABLE {t_idx}:\n")
                for r_idx, row in enumerate(table.rows):
                    cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                    f.write(f"  R{r_idx}: " + " | ".join(cells) + "\n")
    except Exception as e:
        print(f"Failed {filepath}: {e}")

dump_docx('c:/SIAC/SKEL Human Capital 360.docx', 'c:/SIAC/scratch/dump1.txt')
dump_docx('c:/SIAC/SKEL_Human_Capital_360_Product_Bible_v1.docx', 'c:/SIAC/scratch/dump2.txt')
print("Done dumping.")
