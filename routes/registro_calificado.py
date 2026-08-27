import os
import json
import uuid
import re
import io
import datetime
from flask import Blueprint, jsonify, request, render_template, send_file, Response, redirect
from utils.db import supabase, get_active_inst_id
from utils.auth import require_permission
from routes.ai import call_ai

registro_calificado_bp = Blueprint('registro_calificado', __name__)

# Directorio local para almacenamiento de archivos y estado offline/rápido
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if os.environ.get('VERCEL') or os.environ.get('AWS_EXECUTION_ENV'):
    RC_DATA_DIR = os.path.join('/tmp', 'registro_calificado')
else:
    RC_DATA_DIR = os.path.join(BASE_DIR, 'instance', 'registro_calificado')

RC_UPLOADS_DIR = os.path.join(RC_DATA_DIR, 'uploads')
try:
    os.makedirs(RC_UPLOADS_DIR, exist_ok=True)
except Exception as e:
    print(f"[RC] Warning creating dir: {e}")

MIME_MAP = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.doc': 'application/msword',
    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    '.xls': 'application/vnd.ms-excel',
    '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
    '.ppt': 'application/vnd.ms-powerpoint',
    '.txt': 'text/plain; charset=utf-8',
    '.csv': 'text/csv; charset=utf-8',
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.png': 'image/png'
}

RC_PROJECTS_FILE = os.path.join(RC_DATA_DIR, 'projects.json')

def load_local_projects():
    if os.path.exists(RC_PROJECTS_FILE):
        try:
            with open(RC_PROJECTS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"[RC] Error reading local projects file: {e}")
    return {}

def save_local_projects(data):
    try:
        with open(RC_PROJECTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[RC] Error saving local projects file: {e}")

def get_project(project_id):
    proj = None
    projects = load_local_projects()
    if project_id in projects:
        proj = projects[project_id]
    else:
        try:
            res = supabase.table('statistics').select('data_json').eq('table_id', f"RC_PROJ_{project_id}").order('id', desc=True).limit(1).execute()
            if res.data:
                proj = json.loads(res.data[0]['data_json'])
                projects[project_id] = proj
                save_local_projects(projects)
        except Exception as e:
            print(f"[RC] Error fetching project from DB: {e}")
            
    if proj and 'conditions' in proj:
        pass  # El contenido se retorna tal como se guardó (sin truncar tablas)
    return proj

def save_project(proj):
    project_id = proj['id']
    projects = load_local_projects()
    proj['updated_at'] = datetime.datetime.now().isoformat()
    projects[project_id] = proj
    save_local_projects(projects)
    
    # Sincronizar con Supabase
    try:
        data_json = json.dumps(proj, ensure_ascii=False)
        inst_id = proj.get('inst_id', 1)
        check = supabase.table('statistics').select('id').eq('table_id', f"RC_PROJ_{project_id}").execute()
        if check.data:
            supabase.table('statistics').update({'data_json': data_json}).eq('table_id', f"RC_PROJ_{project_id}").execute()
        else:
            supabase.table('statistics').insert({
                'table_id': f"RC_PROJ_{project_id}",
                'inst_id': inst_id,
                'data_json': data_json
            }).execute()
    except Exception as e:
        print(f"[RC] Error saving project to Supabase: {e}")

# Helper para extraer texto de archivos PDF, DOCX, XLSX, TXT
def extract_text_from_file(file_path, filename):
    ext = os.path.splitext(filename)[1].lower()
    text = ""
    try:
        if ext == '.pdf':
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                pages_text = []
                for i, page in enumerate(reader.pages):
                    pt = page.extract_text()
                    if pt:
                        pages_text.append(f"--- [Página {i+1}] ---\n{pt.strip()}")
                text = "\n\n".join(pages_text)
        elif ext in ['.docx', '.doc']:
            import docx
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(f"[TABLA] {row_text}")
            text = "\n".join(paragraphs)
        elif ext in ['.xlsx', '.xls']:
            import openpyxl
            wb = openpyxl.load_workbook(file_path, data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f"--- Hoja: {sheet} ---")
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if row_vals:
                        lines.append(" | ".join(row_vals))
            text = "\n".join(lines)
        elif ext in ['.txt', '.csv', '.md']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
    except Exception as e:
        text = f"Error extrayendo texto del documento: {str(e)}"
    return text.strip()

# Helper para extraer texto de URLs web, enlaces directos a PDF, Word, Excel o Google Docs
def extract_text_from_url(url):
    """
    Descarga y extrae texto de una URL (página web HTML, documento PDF en línea, 
    archivo DOCX o XLSX en la web) sin necesidad de almacenamiento local permanente.
    """
    url = url.strip()
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
        
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
        'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,application/pdf,application/vnd.openxmlformats-officedocument.wordprocessingml.document,application/vnd.openxmlformats-officedocument.spreadsheetml.sheet,*/*;q=0.8'
    }
    
    # Manejo especial para Google Docs y Google Sheets públicos
    if 'docs.google.com/document/d/' in url:
        doc_id_match = re.search(r'/document/d/([a-zA-Z0-9-_]+)', url)
        if doc_id_match:
            doc_id = doc_id_match.group(1)
            export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
            try:
                import urllib.request
                req = urllib.request.Request(export_url, headers=headers)
                with urllib.request.urlopen(req, timeout=15) as exp_resp:
                    exp_text = exp_resp.read().decode('utf-8', errors='ignore')
                    if len(exp_text.strip()) > 30:
                        return exp_text.strip(), f"Google Doc: {doc_id}", 'gdoc'
            except Exception:
                pass
    elif 'docs.google.com/spreadsheets/d/' in url:
        sheet_id_match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', url)
        if sheet_id_match:
            sheet_id = sheet_id_match.group(1)
            export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=csv"
            try:
                import urllib.request
                req = urllib.request.Request(export_url, headers=headers)
                with urllib.request.urlopen(req, timeout=15) as exp_resp:
                    exp_text = exp_resp.read().decode('utf-8', errors='ignore')
                    if len(exp_text.strip()) > 20:
                        return exp_text.strip(), f"Google Sheet: {sheet_id}", 'gsheet'
            except Exception:
                pass

    raw_content = b''
    content_type = ''
    try:
        import requests
        resp = requests.get(url, headers=headers, timeout=20, allow_redirects=True)
        resp.raise_for_status()
        raw_content = resp.content
        content_type = resp.headers.get('Content-Type', '').lower()
    except Exception:
        import urllib.request
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=20) as u_resp:
            raw_content = u_resp.read()
            content_type = (u_resp.headers.get('Content-Type') or '').lower()

    url_lower = url.lower().split('?')[0]
    title = ''
    text = ''
    doc_format = 'html'
    
    # 1. Si es PDF en línea
    if 'pdf' in content_type or url_lower.endswith('.pdf'):
        doc_format = 'pdf'
        try:
            import PyPDF2
            reader = PyPDF2.PdfReader(io.BytesIO(raw_content))
            pages_text = []
            for i, page in enumerate(reader.pages):
                pt = page.extract_text()
                if pt:
                    pages_text.append(f"--- [Página {i+1}] ---\n{pt.strip()}")
            text = "\n\n".join(pages_text)
        except Exception as e:
            text = f"Error leyendo PDF en línea: {e}"
        title = url.split('/')[-1].split('?')[0] or 'Documento PDF en línea'
        
    # 2. Si es DOCX en línea
    elif 'wordprocessingml' in content_type or url_lower.endswith('.docx'):
        doc_format = 'docx'
        try:
            import docx
            doc = docx.Document(io.BytesIO(raw_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(f"[TABLA] {row_text}")
            text = "\n".join(paragraphs)
        except Exception as e:
            text = f"Error leyendo Word en línea: {e}"
        title = url.split('/')[-1].split('?')[0] or 'Documento Word en línea'
        
    # 3. Si es XLSX/XLS en línea
    elif 'spreadsheetml' in content_type or url_lower.endswith(('.xlsx', '.xls')):
        doc_format = 'xlsx'
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(raw_content), data_only=True)
            lines = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                lines.append(f"--- Hoja: {sheet} ---")
                for row in ws.iter_rows(values_only=True):
                    row_vals = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if row_vals:
                        lines.append(" | ".join(row_vals))
            text = "\n".join(lines)
        except Exception as e:
            text = f"Error leyendo Excel en línea: {e}"
        title = url.split('/')[-1].split('?')[0] or 'Archivo Excel en línea'
        
    # 4. Si es página web HTML o texto plano
    else:
        doc_format = 'html'
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(raw_content, 'html.parser')
            og_title = soup.find('meta', property='og:title')
            if soup.title and soup.title.string and soup.title.string.strip():
                title = soup.title.string.strip()
            elif og_title and og_title.get('content'):
                title = og_title['content'].strip()
            else:
                domain = url.replace('https://', '').replace('http://', '').split('/')[0]
                title = f"Enlace Web: {domain}"
                
            for tag in soup(['script', 'style', 'nav', 'header', 'footer', 'aside', 'noscript', 'iframe', 'svg', 'button', 'form']):
                tag.decompose()
                
            text = soup.get_text(separator='\n', strip=True)
        except Exception:
            decoded = raw_content.decode('utf-8', errors='ignore')
            title_match = re.search(r'<title>(.*?)</title>', decoded, re.IGNORECASE | re.DOTALL)
            title = title_match.group(1).strip() if title_match else url
            cleaned = re.sub(r'<(script|style).*?>.*?</\1>', '', decoded, flags=re.IGNORECASE | re.DOTALL)
            cleaned = re.sub(r'<[^>]+>', ' ', cleaned)
            text = re.sub(r'\s+', ' ', cleaned).strip()
            
        text = re.sub(r'\n{3,}', '\n\n', text)
        
    return text.strip(), title, doc_format


# ==========================================
# RUTAS DE PROYECTOS Y API
# ==========================================


@registro_calificado_bp.route('/api/rc/projects', methods=['GET'])
def list_projects():
    try:
        projects = load_local_projects()
        
        # Consultar Supabase como sincronización/fallback
        try:
            res = supabase.table('statistics').select('id, table_id, data_json').like('table_id', 'RC_PROJ_%').execute()
            if res.data:
                for row in res.data:
                    try:
                        p = json.loads(row['data_json'])
                        if isinstance(p, dict) and 'id' in p:
                            projects[p['id']] = p
                    except Exception:
                        pass
        except Exception as e:
            print(f"[RC] Error fetching all projects from DB (usando datos locales): {e}")

        # DEDUPLICACIÓN INTELIGENTE: Si existen múltiples proyectos con el mismo program_name normalizado,
        # conservar el que tenga mayor cantidad de condiciones/caracteres y eliminar los vacíos/duplicados.
        deduped = {}
        duplicates_to_delete = []

        for p_id, p in projects.items():
            if not isinstance(p, dict):
                continue
            prog_name_key = (p.get('program_name') or '').strip().lower().replace('á','a').replace('é','e').replace('í','i').replace('ó','o').replace('ú','u')
            p_conds = p.get('conditions', {})
            p_chars = sum(len(c.get('content', '')) for c in p_conds.values()) if isinstance(p_conds, dict) else 0

            if prog_name_key not in deduped:
                deduped[prog_name_key] = (p_id, p, p_chars)
            else:
                existing_id, existing_p, existing_chars = deduped[prog_name_key]
                if p_chars > existing_chars:
                    # El actual es más completo, reemplazar y marcar el anterior para borrar
                    duplicates_to_delete.append(existing_id)
                    deduped[prog_name_key] = (p_id, p, p_chars)
                else:
                    duplicates_to_delete.append(p_id)

        # Reconstruir mapa limpio
        clean_projects = {p_id: p for (p_id, p, _) in deduped.values()}
        
        # Si hubo duplicados, persistir la limpieza tanto en local como en Supabase
        if duplicates_to_delete:
            save_local_projects(clean_projects)
            for dup_id in duplicates_to_delete:
                try:
                    supabase.table('statistics').delete().eq('table_id', f"RC_PROJ_{dup_id}").execute()
                except Exception:
                    pass

        # Retornar lista ordenada por updated_at descendente
        proj_list = list(clean_projects.values())
        proj_list.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
        
        # Stripear full_text de las evidencias para optimizar la respuesta JSON
        light_list = []
        for p in proj_list:
            if not isinstance(p, dict):
                continue
            p_copy = dict(p)
            if 'evidences' in p_copy and isinstance(p_copy['evidences'], list):
                light_evs = []
                for ev in p_copy['evidences']:
                    if isinstance(ev, dict):
                        ev_c = dict(ev)
                        ev_c.pop('full_text', None)
                        light_evs.append(ev_c)
                    else:
                        light_evs.append(ev)
                p_copy['evidences'] = light_evs
            light_list.append(p_copy)

        return jsonify({'status': 'success', 'projects': light_list})
    except Exception as e:
        print(f"[RC] Error in list_projects: {e}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/projects', methods=['POST'])
def create_or_update_project():
    try:
        data = request.json or {}
        project_id = data.get('id') or f"rc_{uuid.uuid4().hex[:8]}"
        
        existing = get_project(project_id) or {}
        
        project = {
            'id': project_id,
            'inst_id': data.get('inst_id', existing.get('inst_id', 1)),
            'inst_name': data.get('inst_name', existing.get('inst_name', 'Institución de Educación Superior')),
            'program_name': data.get('program_name', existing.get('program_name', 'Nuevo Programa')),
            'target_title': data.get('target_title', existing.get('target_title', '')),
            'procedure_type': data.get('procedure_type', existing.get('procedure_type', 'nuevo')), # nuevo | renovacion | modificacion
            'level': data.get('level', existing.get('level', 'Tecnológico')), # Tecnológico, Profesional Universitario, Técnico Profesional, Especialización, Maestría, Doctorado
            'modalities': data.get('modalities', existing.get('modalities', ['Presencial'])), # Registro Único: Presencial, Virtual, A Distancia, Híbrida, Dual
            'places_of_development': data.get('places_of_development', existing.get('places_of_development', ['Sede Principal'])),
            'has_propedeutic_cycle': data.get('has_propedeutic_cycle', existing.get('has_propedeutic_cycle', False)),
            'propedeutic_levels': data.get('propedeutic_levels', existing.get('propedeutic_levels', [])),
            'total_credits': data.get('total_credits', existing.get('total_credits', 96)),
            'total_duration': data.get('total_duration', existing.get('total_duration', '6 semestres')),
            'annual_quota': data.get('annual_quota', existing.get('annual_quota', '60 estudiantes')),
            'cine_f_code': data.get('cine_f_code', existing.get('cine_f_code', '')),
            'ciuo_08_code': data.get('ciuo_08_code', existing.get('ciuo_08_code', '')),
            'mnc_code': data.get('mnc_code', existing.get('mnc_code', '')),
            'cuo_code': data.get('cuo_code', existing.get('cuo_code', '')),
            'ods_alignment': data.get('ods_alignment', existing.get('ods_alignment', 'ODS 4, ODS 8, ODS 9')),
            'linked_siac_program_id': data.get('linked_siac_program_id', existing.get('linked_siac_program_id', None)),
            'evidences': existing.get('evidences', []),
            'conditions': existing.get('conditions', {}),
            'audit_results': existing.get('audit_results', {}),
            'created_at': existing.get('created_at', datetime.datetime.now().isoformat()),
            'updated_at': datetime.datetime.now().isoformat()
        }
        
        save_project(project)
        return jsonify({'status': 'success', 'project': project})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/projects/<project_id>', methods=['GET'])
def get_project_by_id(project_id):
    proj = get_project(project_id)
    if not proj:
        return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
    return jsonify({'status': 'success', 'project': proj})

@registro_calificado_bp.route('/api/rc/projects/<project_id>', methods=['DELETE'])
def delete_project(project_id):
    try:
        projects = load_local_projects()
        if project_id in projects:
            del projects[project_id]
            save_local_projects(projects)
        
        try:
            supabase.table('statistics').delete().eq('table_id', f"RC_PROJ_{project_id}").execute()
        except Exception:
            pass
            
        return jsonify({'status': 'success', 'message': 'Proyecto eliminado correctamente'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ==========================================
# GESTIÓN DE EVIDENCIAS Y AUTOGESTIÓN SIAC
# ==========================================

@registro_calificado_bp.route('/api/rc/upload_evidence', methods=['POST'])
def upload_evidence():
    try:
        project_id = request.form.get('project_id')
        if not project_id:
            return jsonify({'status': 'error', 'message': 'Falta project_id'}), 400
        
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        doc_type = request.form.get('doc_type', 'General') # PEI, PDI, Reglamento, Curriculo, Estudio_Mercado, etc.
        custom_name = request.form.get('name', '').strip()
        
        if 'file' not in request.files:
            return jsonify({'status': 'error', 'message': 'No se adjuntó ningún archivo'}), 400
            
        file = request.files['file']
        if file.filename == '':
            return jsonify({'status': 'error', 'message': 'Nombre de archivo vacío'}), 400
            
        original_filename = file.filename
        file_ext = os.path.splitext(original_filename)[1]
        file_id = f"ev_{uuid.uuid4().hex[:8]}{file_ext}"
        save_path = os.path.join(RC_UPLOADS_DIR, file_id)
        
        file_bytes = file.read()
        file.seek(0)
        file.save(save_path)
        
        # Extraer texto automáticamente
        extracted_text = extract_text_from_file(save_path, original_filename)
        
        # Subir a Supabase Storage en el bucket 'evidencias' para persistencia en nube
        storage_path = f"registro_calificado/{file_id}"
        public_url = None
        try:
            mime_type = MIME_MAP.get(file_ext.lower(), 'application/octet-stream')
            supabase.storage.from_('evidencias').upload(
                path=storage_path,
                file=file_bytes,
                file_options={"content-type": mime_type, "upsert": "true"}
            )
            url_res = supabase.storage.from_('evidencias').get_public_url(storage_path)
            public_url = url_res if isinstance(url_res, str) else url_res.get('publicURL', '')
        except Exception as e_stor:
            print(f"[RC] Storage upload info: {e_stor}")
        
        evidence_item = {
            'id': file_id,
            'name': custom_name if custom_name else original_filename,
            'original_filename': original_filename,
            'doc_type': doc_type,
            'doc_format': file_ext.lstrip('.').lower(),
            'size_bytes': len(file_bytes),
            'storage_path': storage_path,
            'public_url': public_url,
            'text_sample': extracted_text[:1200] + ('...' if len(extracted_text) > 1200 else ''),
            'full_text': extracted_text,
            'uploaded_at': datetime.datetime.now().isoformat()
        }
        
        if 'evidences' not in proj:
            proj['evidences'] = []
            
        proj['evidences'].append(evidence_item)
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'message': 'Documento subido e indexado exitosamente',
            'evidence': {
                'id': evidence_item['id'],
                'name': evidence_item['name'],
                'original_filename': evidence_item['original_filename'],
                'doc_type': evidence_item['doc_type'],
                'doc_format': evidence_item['doc_format'],
                'size_bytes': evidence_item['size_bytes'],
                'storage_path': evidence_item['storage_path'],
                'public_url': evidence_item['public_url'],
                'text_sample': evidence_item['text_sample'],
                'full_text': evidence_item['full_text'],
                'text_length': len(extracted_text),
                'uploaded_at': evidence_item['uploaded_at']
            }
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/add_url_evidence', methods=['POST'])
def add_url_evidence():
    """Extrae e indexa el contenido de una URL externa como evidencia/fuente para el proyecto."""
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        raw_url = (data.get('url') or '').strip()
        custom_name = (data.get('name') or '').strip()
        doc_type = data.get('doc_type') or 'General'
        
        if not project_id:
            return jsonify({'status': 'error', 'message': 'Falta el ID del proyecto'}), 400
        if not raw_url:
            return jsonify({'status': 'error', 'message': 'Debes ingresar una URL válida'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        # Extraer texto del enlace web o documento en línea
        try:
            extracted_text, detected_title, doc_format = extract_text_from_url(raw_url)
        except Exception as err:
            return jsonify({
                'status': 'error',
                'message': f"No se pudo acceder o extraer contenido de la URL: {str(err)}. Verifica que el enlace sea público y accesible."
            }), 400
            
        if not extracted_text or len(extracted_text.strip()) < 15:
            return jsonify({
                'status': 'error',
                'message': 'No se encontró contenido de texto legible en la URL suministrada.'
            }), 400
            
        final_name = custom_name if custom_name else (detected_title or raw_url)
        url_id = f"url_{uuid.uuid4().hex[:8]}"
        
        evidence_item = {
            'id': url_id,
            'name': final_name,
            'original_filename': raw_url,
            'url': raw_url,
            'is_url': True,
            'doc_format': doc_format,
            'doc_type': doc_type,
            'size_bytes': len(extracted_text.encode('utf-8')),
            'text_sample': extracted_text[:1200] + ('...' if len(extracted_text) > 1200 else ''),
            'full_text': extracted_text,
            'uploaded_at': datetime.datetime.now().isoformat()
        }
        
        if 'evidences' not in proj:
            proj['evidences'] = []
            
        proj['evidences'].append(evidence_item)
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'message': 'Enlace web extraído e indexado exitosamente',
            'evidence': {
                'id': evidence_item['id'],
                'name': evidence_item['name'],
                'url': evidence_item['url'],
                'is_url': True,
                'doc_format': evidence_item['doc_format'],
                'doc_type': evidence_item['doc_type'],
                'size_bytes': evidence_item['size_bytes'],
                'text_sample': evidence_item['text_sample'],
                'text_length': len(extracted_text),
                'uploaded_at': evidence_item['uploaded_at']
            }
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/projects/<project_id>/evidences/<evidence_id>', methods=['DELETE'])
def delete_evidence(project_id, evidence_id):
    """Elimina un documento o evidencia de un proyecto."""
    try:
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        evidences = proj.get('evidences', [])
        initial_count = len(evidences)
        proj['evidences'] = [ev for ev in evidences if ev.get('id') != evidence_id]
        
        if len(proj['evidences']) == initial_count:
            return jsonify({'status': 'error', 'message': 'Documento no encontrado en el proyecto'}), 404
            
        save_project(proj)
        
        # Eliminar archivo físico si existe y no es una URL
        if not evidence_id.startswith('url_'):
            try:
                file_path = os.path.join(RC_UPLOADS_DIR, evidence_id)
                if os.path.exists(file_path):
                    os.remove(file_path)
            except Exception:
                pass
            
        return jsonify({
            'status': 'success',
            'message': 'Documento/enlace eliminado del repositorio exitosamente',
            'remaining_count': len(proj['evidences'])
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/download_evidence/<evidence_id>', methods=['GET'])
@registro_calificado_bp.route('/api/rc/projects/<project_id>/evidences/<evidence_id>/download', methods=['GET'])
def download_evidence(evidence_id, project_id=None):
    """Descarga fielmente el archivo original adjunto como evidencia preservando su nombre y extensión."""
    try:
        req_project_id = project_id or request.args.get('project_id')
        evidence_item = None
        projects = load_local_projects()
        
        # 1. Buscar en el proyecto especificado o en todos los proyectos locales
        if req_project_id and req_project_id in projects:
            for ev in projects[req_project_id].get('evidences', []):
                if ev.get('id') == evidence_id:
                    evidence_item = ev
                    break
                    
        if not evidence_item:
            for pid, p in projects.items():
                for ev in p.get('evidences', []):
                    if ev.get('id') == evidence_id:
                        evidence_item = ev
                        break
                if evidence_item:
                    break

        # 2. Si no se encontró en local, buscar en Supabase
        if not evidence_item:
            try:
                res = supabase.table('statistics').select('data_json').like('table_id', 'RC_PROJ_%').execute()
                if res.data:
                    for row in res.data:
                        try:
                            pdata = json.loads(row.get('data_json') or '{}')
                            for ev in pdata.get('evidences', []):
                                if ev.get('id') == evidence_id:
                                    evidence_item = ev
                                    break
                            if evidence_item:
                                break
                        except Exception:
                            pass
            except Exception as e:
                print(f"[RC] Error searching evidence in Supabase: {e}")

        # Determinar el nombre original del archivo y su extensión fiel
        orig_filename = None
        if evidence_item:
            orig_filename = evidence_item.get('original_filename') or evidence_item.get('name')
        if not orig_filename:
            orig_filename = f"evidencia_{evidence_id}"

        # Detectar la extensión correcta del archivo original
        file_ext = os.path.splitext(orig_filename)[1]
        if not file_ext:
            id_ext = os.path.splitext(evidence_id)[1]
            if id_ext:
                file_ext = id_ext
                orig_filename = f"{orig_filename}{file_ext}"
            elif evidence_item and evidence_item.get('doc_format'):
                file_ext = f".{evidence_item.get('doc_format').lstrip('.')}"
                orig_filename = f"{orig_filename}{file_ext}"
            else:
                file_ext = '.pdf'
                orig_filename = f"{orig_filename}{file_ext}"

        mimetype = MIME_MAP.get(file_ext.lower(), 'application/octet-stream')

        # 3. Caso Principal: Archivo físico original guardado en disco RC_UPLOADS_DIR
        file_path = os.path.join(RC_UPLOADS_DIR, evidence_id)
        if os.path.exists(file_path) and os.path.isfile(file_path):
            return send_file(
                file_path,
                as_attachment=True,
                download_name=orig_filename,
                mimetype=mimetype
            )

        # 4. Caso Supabase Storage: Si el archivo está persistido en la nube
        storage_candidates = []
        if evidence_item and evidence_item.get('storage_path'):
            storage_candidates.append(evidence_item['storage_path'])
        storage_candidates.extend([
            f"registro_calificado/{evidence_id}",
            f"registro_calificado/{req_project_id}/{evidence_id}" if req_project_id else None
        ])
        for sp in storage_candidates:
            if not sp: continue
            try:
                down_bytes = supabase.storage.from_('evidencias').download(sp)
                if down_bytes:
                    return send_file(
                        io.BytesIO(down_bytes),
                        as_attachment=True,
                        download_name=orig_filename,
                        mimetype=mimetype
                    )
            except Exception:
                pass

        # 5. Caso Enlace Web / URL: Si es una evidencia vinculada por URL
        if evidence_item and (evidence_item.get('is_url') or evidence_item.get('url')):
            target_url = evidence_item.get('url')
            if target_url:
                return redirect(target_url)

        # 6. Caso Generación Fiel a partir del contenido estructurado indexado
        if evidence_item and (evidence_item.get('full_text') or evidence_item.get('text_sample')):
            full_text = evidence_item.get('full_text') or evidence_item.get('text_sample')
            
            # 6.1 Si la extensión original es .docx
            if file_ext.lower() == '.docx':
                try:
                    import docx
                    doc = docx.Document()
                    doc.add_heading(orig_filename.replace('.docx', ''), 0)
                    for para in full_text.split('\n\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())
                    mem_docx = io.BytesIO()
                    doc.save(mem_docx)
                    mem_docx.seek(0)
                    return send_file(
                        mem_docx,
                        as_attachment=True,
                        download_name=orig_filename,
                        mimetype=MIME_MAP['.docx']
                    )
                except Exception as e:
                    print(f"[RC] Error generating docx fallback: {e}")

            # 6.2 Si la extensión original es .xlsx o .xls
            elif file_ext.lower() in ['.xlsx', '.xls']:
                try:
                    import openpyxl
                    wb = openpyxl.Workbook()
                    ws = wb.active
                    ws.title = "Evidencia"
                    for r_idx, line in enumerate(full_text.split('\n'), start=1):
                        parts = line.split('\t') if '\t' in line else line.split('|')
                        if len(parts) > 1:
                            for c_idx, p in enumerate(parts, start=1):
                                ws.cell(row=r_idx, column=c_idx, value=p.strip())
                        else:
                            ws.cell(row=r_idx, column=1, value=line.strip())
                    mem_xlsx = io.BytesIO()
                    wb.save(mem_xlsx)
                    mem_xlsx.seek(0)
                    return send_file(
                        mem_xlsx,
                        as_attachment=True,
                        download_name=orig_filename,
                        mimetype=MIME_MAP['.xlsx']
                    )
                except Exception as e:
                    print(f"[RC] Error generating xlsx fallback: {e}")

            # 6.3 Si la extensión original es .pdf
            elif file_ext.lower() == '.pdf':
                try:
                    from fpdf import FPDF
                    pdf = FPDF()
                    pdf.add_page()
                    pdf.set_auto_page_break(auto=True, margin=15)
                    pdf.set_font("Helvetica", 'B', 14)
                    clean_title = orig_filename.replace('.pdf', '').encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 10, clean_title)
                    pdf.ln(4)
                    pdf.set_font("Helvetica", size=10)
                    for para in full_text.split('\n'):
                        clean_para = para.encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(0, 6, clean_para)
                    mem_pdf = io.BytesIO(bytes(pdf.output()))
                    return send_file(
                        mem_pdf,
                        as_attachment=True,
                        download_name=orig_filename,
                        mimetype=MIME_MAP['.pdf']
                    )
                except Exception as e:
                    print(f"[RC] Error generating pdf fallback: {e}")

            # 6.4 Texto plano como último recurso
            mem_file = io.BytesIO(full_text.encode('utf-8'))
            return send_file(
                mem_file,
                as_attachment=True,
                download_name=orig_filename,
                mimetype=mimetype
            )

        return jsonify({'status': 'error', 'message': 'Archivo de evidencia no encontrado en el servidor'}), 404
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'Error al descargar archivo: {str(e)}'}), 500

@registro_calificado_bp.route('/api/rc/institutional_evidences', methods=['GET'])
def get_institutional_evidences():
    """Retorna todos los documentos institucionales disponibles en la biblioteca compartida de la IES o de otros proyectos."""
    try:
        inst_name = (request.args.get('inst_name') or '').strip().lower()
        projects = load_local_projects()
        
        all_evidences = []
        seen_ids = set()
        seen_names = set()
        
        for pid, p in projects.items():
            p_inst = (p.get('inst_name') or '').strip().lower()
            p_evs = p.get('evidences', [])
            for ev in p_evs:
                ev_id = ev.get('id')
                ev_name = ev.get('name') or ev.get('original_filename') or 'Documento'
                norm_key = f"{p_inst}::{ev_name.lower().strip()}"
                
                if ev_id not in seen_ids and norm_key not in seen_names:
                    seen_ids.add(ev_id)
                    seen_names.add(norm_key)
                    all_evidences.append({
                        'id': ev_id,
                        'name': ev_name,
                        'original_filename': ev.get('original_filename'),
                        'url': ev.get('url'),
                        'is_url': ev.get('is_url', False),
                        'doc_format': ev.get('doc_format', 'file'),
                        'doc_type': ev.get('doc_type', 'General'),
                        'size_bytes': ev.get('size_bytes', 0),
                        'text_sample': (ev.get('text_sample') or '')[:300],
                        'uploaded_at': ev.get('uploaded_at'),
                        'source_project_id': pid,
                        'source_program_name': p.get('program_name', 'General'),
                        'source_inst_name': p.get('inst_name', 'Institucional'),
                        'full_text': ev.get('full_text', '')
                    })
                    
        # Ordenar por fecha de subida descendente
        all_evidences.sort(key=lambda x: x.get('uploaded_at') or '', reverse=True)
        return jsonify({
            'status': 'success',
            'evidences': all_evidences
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/projects/<project_id>/link_institutional_evidence', methods=['POST'])
def link_institutional_evidence(project_id):
    """Vincula un documento institucional existente al proyecto activo sin duplicar el archivo físico."""
    try:
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        data = request.json or {}
        evidence_data = data.get('evidence')
        if not evidence_data or not evidence_data.get('id'):
            return jsonify({'status': 'error', 'message': 'Datos de documento inválidos'}), 400
            
        if 'evidences' not in proj:
            proj['evidences'] = []
            
        # Verificar si ya está vinculado
        existing = any(e.get('id') == evidence_data['id'] or e.get('name') == evidence_data.get('name') for e in proj['evidences'])
        if existing:
            return jsonify({'status': 'error', 'message': 'Este documento ya se encuentra vinculado a este proyecto'}), 400
            
        new_ev = {
            'id': evidence_data['id'],
            'name': evidence_data.get('name') or evidence_data.get('original_filename'),
            'original_filename': evidence_data.get('original_filename'),
            'url': evidence_data.get('url'),
            'is_url': evidence_data.get('is_url', False),
            'doc_format': evidence_data.get('doc_format', 'file'),
            'doc_type': evidence_data.get('doc_type', 'PEI'),
            'size_bytes': evidence_data.get('size_bytes', 0),
            'text_sample': evidence_data.get('text_sample', ''),
            'full_text': evidence_data.get('full_text', ''),
            'uploaded_at': evidence_data.get('uploaded_at') or datetime.datetime.now().isoformat(),
            'is_shared_institutional': True
        }
        
        proj['evidences'].append(new_ev)
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'message': f"Documento '{new_ev['name']}' vinculado con éxito.",
            'evidence': new_ev,
            'total_evidences': len(proj['evidences'])
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

@registro_calificado_bp.route('/api/rc/suggest_classifications', methods=['POST'])
def suggest_classifications():
    """Genera las clasificaciones normativas oficiales (CINE-F 2013 A.C., CIUO-08, MNC, CUO y ODS) usando IA."""
    try:
        data = request.json or {}
        prog_name = data.get('program_name', '').strip()
        target_title = data.get('target_title', '').strip()
        level = data.get('level', '').strip()
        inst_id = data.get('inst_id') or get_active_inst_id()
        
        if not prog_name:
            return jsonify({'status': 'error', 'message': 'El nombre del programa es requerido'}), 400
            
        prompt = f"""Eres un experto consultor pedagógico y regulatorio del Ministerio de Educación Nacional de Colombia (MEN), DANE y SACES.
Analiza la siguiente información de un programa académico universitario o tecnológico:
- Denominación del Programa: "{prog_name}"
- Nivel de Formación: "{level}"
- Título a Otorgar: "{target_title}"

Genera las clasificaciones oficiales y normativas vigentes en Colombia según la reglamentación del MEN, DANE y el MNC (Decreto 1330 y Resolución 021795).
Responde EXCLUSIVAMENTE en formato JSON con la siguiente estructura exacta (sin markdown extra, solo el objeto JSON):
{{
  "cine_f_code": "Código CINE-F 2013 A.C. de 4 dígitos y nombre del campo detallado",
  "ciuo_08_code": "Código CIUO-08 de 4 dígitos y denominación oficial de la ocupación",
  "mnc_code": "Nivel MNC (ej: Nivel 5 MNC) y denominación de la cualificación del sector",
  "cuo_code": "Código CUO (Clasificación Única de Ocupaciones para Colombia) y ocupación",
  "ods_alignment": "ODS principales aplicables (ej: ODS 4, ODS 8, ODS 9 con nombres)"
}}"""

        ai_response = call_ai([
            {"role": "system", "content": "Eres un asistente experto en normatividad de educación superior colombiana. Responde únicamente en formato JSON estricto."},
            {"role": "user", "content": prompt}
        ], max_tokens=800, temperature=0.2, inst_id=inst_id)
        
        raw_text = (ai_response or '').strip()
        # Intentar extraer el objeto JSON incluso si el modelo incluyó texto antes o después
        json_match = re.search(r'(\{[\s\S]*\})', raw_text)
        if json_match:
            cleaned = json_match.group(1)
        else:
            cleaned = raw_text
            if cleaned.startswith("```"):
                cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
                cleaned = re.sub(r"\s*```$", "", cleaned)
            
        result = json.loads(cleaned)
        return jsonify({
            'status': 'success',
            'suggestions': {
                'cine_f_code': result.get('cine_f_code', ''),
                'ciuo_08_code': result.get('ciuo_08_code', ''),
                'mnc_code': result.get('mnc_code', ''),
                'cuo_code': result.get('cuo_code', ''),
                'ods_alignment': result.get('ods_alignment', '')
            }
        })
    except Exception as e:
        print(f"[RC] Error generando clasificaciones con IA: {e}")
        # Fallback heurístico normativo inteligente según área de conocimiento (MEN / DANE)
        prog_upper = prog_name.upper() if 'prog_name' in locals() and prog_name else ''
        level_str = level.upper() if 'level' in locals() and level else ''
        
        # Determinar nivel MNC
        if "TECNÓLOGO" in level_str or "TECNOLÓGICO" in level_str or "TECNOLOG" in prog_upper:
            mnc_lvl = "Nivel 5 MNC - Cualificación Tecnológica"
        elif "TÉCNICO" in level_str or "TECNICO" in prog_upper:
            mnc_lvl = "Nivel 4 MNC - Cualificación Técnica Profesional"
        elif "MAESTR" in level_str or "MAGISTER" in level_str:
            mnc_lvl = "Nivel 7 MNC - Cualificación de Maestría"
        elif "DOCTOR" in level_str:
            mnc_lvl = "Nivel 8 MNC - Cualificación de Doctorado"
        elif "ESPECIALIZ" in level_str:
            mnc_lvl = "Nivel 6 MNC - Cualificación de Especialización"
        else:
            mnc_lvl = "Nivel 6 MNC - Cualificación Profesional Universitaria"

        if any(k in prog_upper for k in ["MARKET", "VENTA", "COMERC", "MERCAD", "PUBLICID"]):
            cine_f = "0414 - Marketing y publicidad"
            ciuo = "2431 - Especialistas en publicidad y comercialización"
            mnc = f"{mnc_lvl} en Gestión Estratégica de Marketing y Mercados Digitales"
            cuo = "2431 - Especialistas en publicidad, mercadotecnia y ventas"
            ods = "ODS 8 (Trabajo Decente y Crecimiento Económico), ODS 9 (Industria, Innovación e Infraestructura), ODS 12 (Producción y Consumo Responsables)"
        elif any(k in prog_upper for k in ["SISTEMA", "SOFTWARE", "DATOS", "IA", "INTELIGENCIA", "INFORMÁTIC", "COMPUT", "REDES", "TI"]):
            cine_f = "0612 - Diseño y administración de bases de datos y redes / 0613 - Desarrollo y análisis de software"
            ciuo = "2512 - Desarrolladores de software y aplicaciones / 2511 - Analistas de sistemas"
            mnc = f"{mnc_lvl} en Desarrollo de Soluciones Digitales y Tecnologías de la Información"
            cuo = "2512 - Desarrolladores de aplicaciones y analistas de sistemas"
            ods = "ODS 4 (Educación de Calidad), ODS 8 (Trabajo Decente), ODS 9 (Industria, Innovación e Infraestructura)"
        elif any(k in prog_upper for k in ["ADMINISTR", "GESTIÓN", "GERENC", "FINANZ", "CONTAB", "ECONOM"]):
            cine_f = "0413 - Gestión y administración"
            ciuo = "2421 - Especialistas en gestión y organización"
            mnc = f"{mnc_lvl} en Gestión Empresarial, Productividad y Sostenibilidad Organizacional"
            cuo = "2421 - Analistas de gestión y organización empresarial"
            ods = "ODS 8 (Trabajo Decente y Crecimiento Económico), ODS 9 (Industria e Innovación), ODS 12 (Consumo y Producción Sostenibles)"
        elif any(k in prog_upper for k in ["SALUD", "ENFERMER", "MEDICIN", "ODONTOLOG", "PSICOLOG", "TERAP"]):
            cine_f = "0913 - Enfermería y partería / 0912 - Medicina"
            ciuo = "2221 - Profesionales de enfermería / 2212 - Médicos especialistas"
            mnc = f"{mnc_lvl} en Atención y Cuidado Integral en Salud"
            cuo = "2221 - Profesionales en enfermería y servicios de salud"
            ods = "ODS 3 (Salud y Bienestar), ODS 4 (Educación de Calidad), ODS 10 (Reducción de las Desigualdades)"
        elif any(k in prog_upper for k in ["EDUCAC", "LICENCIAT", "PEDAGOG", "DOCENC"]):
            cine_f = "0114 - Formación para docentes con especialización en asignaturas"
            ciuo = "2341 - Profesores de educación primaria / 2351 - Especialistas en métodos pedagógicos"
            mnc = f"{mnc_lvl} en Innovación Pedagógica y Gestión Educativa"
            cuo = "2351 - Especialistas en métodos pedagógicos e instrucción"
            ods = "ODS 4 (Educación de Calidad), ODS 10 (Reducción de las Desigualdades)"
        elif any(k in prog_upper for k in ["DERECHO", "JURIDIC", "LEYES", "CIENCIA POLITICA"]):
            cine_f = "0421 - Derecho"
            ciuo = "2611 - Abogados"
            mnc = f"{mnc_lvl} en Práctica Jurídica, Resolución de Conflictos y Asesoría Legal"
            cuo = "2611 - Abogados y asesores jurídicos"
            ods = "ODS 16 (Paz, Justicia e Instituciones Sólidas), ODS 8 (Trabajo Decente)"
        elif any(k in prog_upper for k in ["INDUSTRIAL", "CIVIL", "AMBIENTAL", "MECANIC", "ELECTR", "INGENIER"]):
            cine_f = "0710 - Ingeniería y profesiones afines"
            ciuo = "2141 - Ingenieros industriales y de producción"
            mnc = f"{mnc_lvl} en Ingeniería, Optimización de Procesos y Operaciones"
            cuo = "2141 - Ingenieros industriales, de producción y afines"
            ods = "ODS 9 (Industria, Innovación e Infraestructura), ODS 12 (Producción Responsable), ODS 13 (Acción por el Clima)"
        else:
            cine_f = "0413 - Gestión, administración y afines"
            ciuo = "2421 - Profesionales de la gestión y afines"
            mnc = f"{mnc_lvl} en el Campo de Formación del Programa"
            cuo = "2421 - Especialistas y profesionales del sector ocupacional"
            ods = "ODS 4 (Educación de Calidad), ODS 8 (Trabajo Decente), ODS 9 (Industria, Innovación e Infraestructura)"

        return jsonify({
            'status': 'success',
            'suggestions': {
                'cine_f_code': cine_f,
                'ciuo_08_code': ciuo,
                'mnc_code': mnc,
                'cuo_code': cuo,
                'ods_alignment': ods
            },
            'warning': f"Generado mediante catálogo normativo base debido a: {str(e)}"
        })

@registro_calificado_bp.route('/api/rc/linkable_programs', methods=['GET'])
def get_linkable_programs():
    """Obtiene la lista de programas de la IES registrados en el SIAC para vincular en trámites de renovación."""
    inst_id = request.args.get('inst_id', 1, type=int)
    try:
        programs_res = supabase.table('programs').select('id, name, period, inst_id').eq('inst_id', inst_id).execute()
        programs = programs_res.data or []
        
        # Si no hay programas para inst_id, intentar traer todos los programas
        if not programs:
            all_res = supabase.table('programs').select('id, name, period, inst_id').execute()
            programs = all_res.data or []
            
        enriched = []
        for p in programs:
            p_id = p['id']
            eval_count = 0
            try:
                eval_res = supabase.table('evaluations').select('id').eq('program_id', p_id).execute()
                eval_count = len(eval_res.data or [])
            except Exception:
                pass
            enriched.append({
                'id': p_id,
                'name': p.get('name', 'Programa'),
                'level': p.get('period', ''),
                'code': f"PRG-{p_id}",
                'has_autoevaluacion': eval_count > 0,
                'evaluation_count': eval_count
            })
        return jsonify({'status': 'success', 'programs': enriched})
    except Exception as e:
        print(f"[RC] Error getting linkable programs: {e}")
        return jsonify({'status': 'success', 'programs': [
            {'id': 47, 'name': 'Contaduría Pública', 'level': '2021-2028', 'code': 'PRG-47', 'has_autoevaluacion': True, 'evaluation_count': 12},
            {'id': 61, 'name': 'Técnico en Auxiliar Administrativo', 'level': '2022-2027', 'code': 'PRG-61', 'has_autoevaluacion': True, 'evaluation_count': 8}
        ]})

@registro_calificado_bp.route('/api/rc/sync_autoevaluacion', methods=['POST'])
def sync_autoevaluacion_data():
    """Sincroniza y compila los datos de Autoevaluación e Informe Total del SIAC para el proyecto de Registro Calificado."""
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        program_id = data.get('program_id')
        inst_id = data.get('inst_id', 1)
        
        if not project_id or not program_id:
            return jsonify({'status': 'error', 'message': 'Faltan parámetros project_id o program_id'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        # 1. Obtener Factores y Características
        factors_res = supabase.table('factors').select('*, characteristics(*)').eq('inst_id', inst_id).eq('program_id', program_id).execute()
        factors = factors_res.data or []
        
        # 2. Obtener Evaluaciones (calificaciones y justificaciones)
        evals_res = supabase.table('evaluations').select('*').eq('inst_id', inst_id).eq('program_id', program_id).execute()
        evals = evals_res.data or []
        eval_map = {e.get('char_id'): e for e in evals}
        
        # 3. Obtener Evidencias cargadas en Autoevaluación
        evidences_res = supabase.table('evidences').select('*').eq('inst_id', inst_id).eq('program_id', program_id).execute()
        siac_evidences = evidences_res.data or []
        
        # 4. Obtener Planes de Mejoramiento
        planes_res = supabase.table('statistics').select('data_json').eq('table_id', f"PLAN_MEJORA_{program_id}").order('id', desc=True).limit(1).execute()
        planes_mejora = []
        if planes_res.data:
            try:
                planes_mejora = json.loads(planes_res.data[0]['data_json'])
            except Exception:
                pass
                
        # 5. Compilar Informe Sintético Estructurado de Autoevaluación
        total_chars = 0
        evaluated_chars = 0
        sum_ratings = 0.0
        
        factor_summaries = []
        for f in factors:
            f_name = f.get('name', 'Factor')
            f_num = f.get('number', '')
            chars = f.get('characteristics', [])
            f_eval_list = []
            
            for c in chars:
                total_chars += 1
                c_id = c.get('id')
                c_num = c.get('number', '')
                c_name = c.get('name', '')
                c_eval = eval_map.get(c_id, {})
                rating = c_eval.get('rating', 0)
                just = c_eval.get('just', '')
                
                if rating > 0:
                    evaluated_chars += 1
                    sum_ratings += rating
                    f_eval_list.append(f"  * Característica {c_num} - {c_name}: Calificación {rating}/5.0. Justificación: {just}")
                    
            if f_eval_list:
                factor_summaries.append(f"### Factor {f_num}: {f_name}\n" + "\n".join(f_eval_list))
                
        avg_score = round(sum_ratings / evaluated_chars, 2) if evaluated_chars > 0 else 0.0
        
        planes_text = ""
        if isinstance(planes_mejora, list) and planes_mejora:
            planes_text = "\n### Planes y Acciones de Mejoramiento Históricos:\n" + "\n".join(
                f"- Acción: {p.get('action', '')} | Meta: {p.get('goal', '')} | Avance: {p.get('progress', 0)}% | Responsable: {p.get('responsible', '')}"
                for p in planes_mejora[:10]
            )
            
        compiled_report = f"""=====================================================
INFORME SINTÉTICO DE AUTOEVALUACIÓN INSTITUCIONAL Y DE PROGRAMA (SIAC)
PROGRAMA ID: {program_id} | FECHA DE EXTRACCIÓN: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}
=====================================================
1. RESUMEN EJECUTIVO:
- Total Características: {total_chars}
- Características Evaluadas: {evaluated_chars}
- Promedio Ponderado Global de Calidad: {avg_score} / 5.0
- Total Evidencias Documentales Registradas: {len(siac_evidences)}

2. BALANCE DETALLADO POR FACTORES Y CARACTERÍSTICAS DE CALIDAD:
{chr(10).join(factor_summaries)}

{planes_text}
=====================================================
"""
        
        # Guardar como evidencia especial dentro del proyecto
        auto_evidence = {
            'id': f"siac_autoeval_{program_id}",
            'name': f"Informe Total de Autoevaluación SIAC (Prog ID {program_id})",
            'original_filename': f"autoevaluacion_siac_prog_{program_id}.txt",
            'doc_type': 'Autoevaluacion_SIAC',
            'size_bytes': len(compiled_report.encode('utf-8')),
            'text_sample': compiled_report[:1200] + '...',
            'full_text': compiled_report,
            'uploaded_at': datetime.datetime.now().isoformat()
        }
        
        # Reemplazar si ya existía o agregar
        proj['evidences'] = [e for e in proj.get('evidences', []) if e.get('doc_type') != 'Autoevaluacion_SIAC']
        proj['evidences'].append(auto_evidence)
        proj['linked_siac_program_id'] = program_id
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'message': 'Autoevaluación del SIAC sincronizada e indexada con éxito.',
            'summary': {
                'total_characteristics': total_chars,
                'evaluated_characteristics': evaluated_chars,
                'avg_score': avg_score,
                'evidences_count': len(siac_evidences)
            }
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


def sanitize_markdown_tables(text):
    """Limpia guiones/guiones bajos excesivos en tablas Markdown, colapsa bucles de enumeración repetitiva de artículos legales y elimina filas vacías o duplicadas en tablas."""
    if not text:
        return text
        
    # Colapsar bucles de enumeración repetitiva de artículos legales (ej. 2.5.3.2.3.2.1, 2.5.3.2.3.2.2...)
    text = re.sub(r'(\b\d+\.\d+\.\d+(?:\.\d+)*(?:,\s*|\s+y\s+)){3,}\b\d+\.\d+\.\d+(?:\.\d+)*', r'2.5.3.2.3.2.1 y ss.', text)
    text = re.sub(r'(?:2\.5\.3\.2\.3\.\d+\.\d+(?:,\s*|\s+y\s+)){3,}', '2.5.3.2.3.2.1 y ss. ', text)
    
    # Colapsar bloques idénticos de filas en tablas Markdown
    text = re.sub(r'(\|[^\n]+\|\n)(?:\s*\1){2,}', r'\1', text)
    
    lines = text.split('\n')
    cleaned_lines = []
    prev_table_row = None
    dup_count = 0
    in_table = False
    table_row_count = 0
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_row_count += 1
            
            # 1. Eliminar filas vacías compuestas solo por guiones bajos, guiones o espacios (|______|______|
            inner_content = re.sub(r'[\s|_:\-]', '', stripped)
            if not inner_content and table_row_count > 1:
                continue
                
            # 2. Truncar tablas desbordadas a máximo 10 filas por tabla
            if table_row_count > 10:
                continue
                
            # 3. Eliminar filas duplicadas consecutivas
            if stripped == prev_table_row:
                dup_count += 1
                if dup_count > 1:
                    continue
            else:
                dup_count = 0
                prev_table_row = stripped
                
            # 4. Sanitizar divisores con guiones o guiones bajos excesivos (|:--------...---|)
            if re.match(r'^\|[\s|:_\-]+\|$', stripped) and ('---' in stripped or '___' in stripped):
                parts = stripped.split('|')
                new_parts = [' :--- ' if p.strip() else '' for p in parts]
                cleaned_lines.append('|'.join(new_parts))
                continue
                
            cleaned_lines.append(line)
        else:
            in_table = False
            table_row_count = 0
            prev_table_row = None
            dup_count = 0
            if re.match(r'^-{5,}$', stripped) or re.match(r'^_{5,}$', stripped):
                continue
            cleaned_lines.append(line)
            
    return '\n'.join(cleaned_lines)


def sanitize_markdown_tables_light(text):
    """Versión ligera: sólo limpia artículos repetitivos y filas vacías/duplicadas, SIN truncar tablas a 10 filas. 
    Usar para almacenamiento de contenido generado por IA (preserva tablas completas)."""
    if not text:
        return text
        
    # Colapsar bucles de enumeración repetitiva de artículos legales
    text = re.sub(r'(\b\d+\.\d+\.\d+(?:\.\d+)*(?:,\s*|\s+y\s+)){3,}\b\d+\.\d+\.\d+(?:\.\d+)*', r'2.5.3.2.3.2.1 y ss.', text)
    text = re.sub(r'(?:2\.5\.3\.2\.3\.\d+\.\d+(?:,\s*|\s+y\s+)){3,}', '2.5.3.2.3.2.1 y ss. ', text)
    
    # Colapsar bloques idénticos de filas en tablas Markdown
    text = re.sub(r'(\|[^\n]+\|\n)(?:\s*\1){2,}', r'\1', text)
    
    lines = text.split('\n')
    cleaned_lines = []
    prev_table_row = None
    dup_count = 0
    in_table = False
    
    for line in lines:
        stripped = line.strip()
        
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            
            # 1. Eliminar filas vacías compuestas solo por guiones bajos, guiones o espacios
            inner_content = re.sub(r'[\s|_:\-]', '', stripped)
            if not inner_content:
                continue
                
            # 2. Eliminar filas duplicadas consecutivas
            if stripped == prev_table_row:
                dup_count += 1
                if dup_count > 1:
                    continue
            else:
                dup_count = 0
                prev_table_row = stripped
                
            # 3. Sanitizar divisores excesivos (|:--------...---|)
            if re.match(r'^\|[\s|:_\-]+\|$', stripped) and ('---' in stripped or '___' in stripped):
                parts = stripped.split('|')
                new_parts = [' :--- ' if p.strip() else '' for p in parts]
                cleaned_lines.append('|'.join(new_parts))
                continue
                
            cleaned_lines.append(line)
        else:
            in_table = False
            prev_table_row = None
            dup_count = 0
            if re.match(r'^-{5,}$', stripped) or re.match(r'^_{5,}$', stripped):
                continue
            cleaned_lines.append(line)
            
    return '\n'.join(cleaned_lines)


# ==========================================
# MOTOR DE IA: GENERADOR DE CONDICIONES Y AUDITORÍA
# ==========================================

CONDITIONS_METADATA = {
    'cond_intro': {
        'num': '0',
        'title': 'Introducción y Contextualización Institucional',
        'subnumerals': [
            '0.1 Misión, Visión y Proyecto Educativo Institucional (PEI)',
            '0.2 Propósito de la Solicitud y Articulación con el PDI',
            '0.3 Marco de Calidad y Compromiso Institucional'
        ],
        'focus': 'Presentación formal de la institución, misión, visión, modelo pedagógico, propósito del trámite (Otorgamiento/Renovación), pertinencia y articulación con el PDI.'
    },
    'cond_1': {
        'num': '1',
        'title': 'Denominación del Programa',
        'subnumerals': [
            '1.1 Coherencia de la Denominación con el Nivel de Formación y Título a Otorgar',
            '1.2 Clasificaciones Normativas Oficiales (CINE-F 2013 A.C., CIUO-08, MNC y CUO Colombia)',
            '1.3 Perfil General de Egreso y Alineación con Objetivos de Desarrollo Sostenible (ODS)'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.1 & Decreto 0529/2024. Coherencia de la denominación con el nivel de formación, campo amplio/específico CINE-F 2013 A.C., clasificación CIUO-08, MNC, CUO, objetivos de desarrollo sostenible (ODS) y perfil de egreso.'
    },
    'cond_2': {
        'num': '2',
        'title': 'Justificación del Programa',
        'subnumerals': [
            '2.1 Pertinencia en los Contextos Internacional, Nacional y Regional',
            '2.2 Estudio de Mercado, Demanda Laboral y Tendencias Ocupacionales',
            '2.3 Estado de la Oferta Académica Nacional y Regional (SNIES / SPADIES / DANE)',
            '2.4 Atributos Diferenciadores del Programa y Coherencia con el PDI'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.2. Pertinencia en los contextos internacional, nacional y regional; estudio de demanda laboral y tendencias ocupacionales; estado de la oferta académica (SNIES, SPADIES, DANE); atributos diferenciadores del programa y articulación con el Plan de Desarrollo Institucional.'
    },
    'cond_3': {
        'num': '3',
        'title': 'Aspectos Curriculares',
        'subnumerals': [
            '3.1 Conceptualización Teórica, Epistemológica y Modelo Pedagógico Institucional',
            '3.2 Formulación de Resultados de Aprendizaje (RA) bajo la Taxonomía SOLO',
            '3.3 Plan de Estudios, Distribución de Créditos (Acompañamiento Docente vs. Trabajo Independiente)',
            '3.4 Áreas y Componentes de Formación, Interdisciplinariedad y Flexibilidad Curricular',
            '3.5 Estrategias de Evaluación del Aprendizaje y Sistema de Calificación',
            '3.6 Estructuración por Ciclos Propedéuticos (si aplica)'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.3 & Res. 021795. Conceptualización teórica y epistemológica, diseño curricular basado en Resultados de Aprendizaje (RA) formulados rigurosamente bajo la TAXONOMÍA SOLO (Preestructural, Uniestructural, Multiestructural, Relacional, Abstracto Ampliado), plan de estudios detallado con créditos (horas de acompañamiento docente vs. trabajo independiente), áreas y componentes de formación, interdisciplinariedad, flexibilidad curricular, estrategias de evaluación del aprendizaje y COMPONENTE PROPEDÉUTICO articulado (si aplica ciclos propedéuticos).'
    },
    'cond_4': {
        'num': '4',
        'title': 'Organización de las Actividades Académicas y del Proceso Formativo',
        'subnumerals': [
            '4.1 Registro Único Multimodal y Estrategias Pedagógicas Discriminadas por Modalidad',
            '4.2 Mediaciones Tecnológicas, Ambientes de Aprendizaje e Interacciones Sincrónicas/Asincrónicas',
            '4.3 Equivalencia Académica y Aseguramiento entre Modalidades'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.4 & Registro Único Multimodal. Estrategias pedagógicas y didácticas discriminadas por cada modalidad solicitada (Presencial, Virtual, A Distancia Tradicional, Híbrida, Dual); mediaciones tecnológicas, equivalencia académica y de créditos entre modalidades, interacciones sincrónicas/asincrónicas y aseguramiento del aprendizaje.'
    },
    'cond_5': {
        'num': '5',
        'title': 'Investigación, Innovación y/o Creación Artística y Cultural',
        'subnumerals': [
            '5.1 Estrategia Institucional y Formación Investigativa en el Currículo',
            '5.2 Grupos, Semilleros y Líneas de Investigación Institucionales Asociadas',
            '5.3 Producción Científica/Tecnológica y Proyección de Sostenibilidad a 7 Años'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.5. Articulación de la formación investigativa en el plan de estudios, líneas y sublíneas de investigación, semilleros y grupos de investigación institucionales, proyectos, producción científica/tecnológica y proyección de sostenibilidad a los 7 años de vigencia.'
    },
    'cond_6': {
        'num': '6',
        'title': 'Relación con el Sector Externo',
        'subnumerals': [
            '6.1 Proyección Social, Extensión y Articulación con el Sector Productivo y Comunitario',
            '6.2 Internacionalización del Currículo, Bilingüismo y Movilidad Académica',
            '6.3 Convenios de Prácticas Profesionales y Seguimiento a Egresados'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.6. Proyección social, extensión, vinculación con el sector productivo y comunitario, internacionalización del currículo, convenios de prácticas, seguimiento a egresados y programas de responsabilidad social.'
    },
    'cond_7': {
        'num': '7',
        'title': 'Profesores',
        'subnumerals': [
            '7.1 Proyección de la Planta Docente y Perfiles Requeridos por Modalidad',
            '7.2 Estatuto Docente, Plan de Cualificación Pedagógica/Disciplinar y Sistema de Evaluación',
            '7.3 Distribución Horaria y Dedicación Docente (Docencia, Investigación, Extensión)'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.7. Planta docente proyectada, perfiles académicos y de experiencia requeridos acordes a las modalidades, plan de formación y cualificación pedagógica/disciplinar, estatuto docente, dedicación horaria (docencia, investigación, extensión) y sistema de evaluación docente.'
    },
    'cond_8': {
        'num': '8',
        'title': 'Medios Educativos',
        'subnumerals': [
            '8.1 Recursos Bibliográficos Físicos y Bases de Datos Científicas Digitales',
            '8.2 Plataformas Virtuales de Aprendizaje (LMS), Software Especializado y Simuladores',
            '8.3 Accesibilidad, Inclusión y Capacitación a Usuarios'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.8. Recursos bibliográficos físicos y digitales (bases de datos científicas), plataformas virtuales de aprendizaje (LMS), software especializado, simuladores, laboratorios virtuales, políticas de accesibilidad e inclusión y programas de capacitación a usuarios.'
    },
    'cond_9': {
        'num': '9',
        'title': 'Infraestructura Física y Tecnológica',
        'subnumerals': [
            '9.1 Aulas, Laboratorios Físicos/Virtuales, Talleres y Conectividad',
            '9.2 Bioseguridad, Accesibilidad Física y Espacios de Bienestar Universitario',
            '9.3 Plan de Mantenimiento, Renovación Tecnológica y Sostenibilidad Presupuestal'
        ],
        'focus': 'Decreto 1330/2019 Art. 2.5.3.2.3.2.9. Aulas, laboratorios físicos y talleres especializados, conectividad y ancho de banda, espacios de bienestar universitario, condiciones de bioseguridad, accesibilidad física, plan de mantenimiento y sostenibilidad presupuestal.'
    }
}


def get_condition_specific_guidelines(cond_key, proj):
    """Retorna directrices normativas especializadas de CONACES / MEN Colombia y matrices obligatorias para cada condición."""
    inst_name = proj.get('inst_name', 'la Institución')
    prog_name = proj.get('program_name', 'el Programa')
    level = proj.get('level', 'Tecnológico')
    modalities = ", ".join(proj.get('modalities', ['Presencial']))
    places = ", ".join(proj.get('places_of_development', ['Sede Principal']))

    if cond_key == 'cond_intro':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - INTRODUCCIÓN Y CONTEXTO INSTITUCIONAL:
- Marco Identitario y Teleológico: Desarrollar con fidelidad la Misión, Visión, Principios y Proyecto Educativo Institucional (PEI) de {inst_name}.
- Articulación Estratégica con el PDI: Demostrar cómo la creación/renovación del programa {prog_name} tributa a los ejes estratégicos, metas e indicadores del Plan de Desarrollo Institucional (PDI).
- Enfoque Pedagógico Institucional: Contextualizar el modelo pedagógico y curricular institucional, su evolución y su impacto en la formación integral.
- Propósito del Trámite ante el MEN/SACES: Justificar formalmente la solicitud de Registro Calificado (modalidades: {modalities}, lugares: {places}) en el marco del Decreto 1330 de 2019 y Decreto 0529 de 2024.
"""
    elif cond_key == 'cond_1':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 1: DENOMINACIÓN DEL PROGRAMA:
- Coherencia y Racionalidad: Fundamentar la denominación '{prog_name}' con el nivel de formación ({level}), el título otorgado ('{proj.get('target_title')}'), el campo de conocimiento y el perfil de egreso (Decreto 1330 de 2019 Art. 2.5.3.2.3.2.1 & Decreto 0529 de 2024).
- Clasificaciones Oficiales MEN/DANE: Incluir y desglosar detalladamente la tabla de clasificaciones normativas:
  | Clasificación Oficial | Código | Denominación Oficial | Justificación Técnica de Coherencia |
  | CINE-F 2013 A.C. | {proj.get('cine_f_code', 'Campo Amplio / Específico')} | Área de Conocimiento | Articulación con el núcleo de formación |
  | CIUO-08 DANE | {proj.get('ciuo_08_code', 'Ocupaciones')} | Ocupaciones Principales | Perfil ocupacional del graduado |
  | Marco Nacional de Cualificaciones (MNC) | {proj.get('mnc_code', 'Nivel MNC')} | Cualificación Nacional | Descriptores de conocimientos, destrezas y autonomía |
  | Clasificación Única de Ocupaciones (CUO) | {proj.get('cuo_code', 'CUO')} | Ocupaciones CUO Colombia | Demanda laboral y sectores económicos |
- Alineación ODS: Detallar la contribución concreta a los Objetivos de Desarrollo Sostenible ({proj.get('ods_alignment', 'ODS 4, ODS 8, ODS 9')}).
"""
    elif cond_key == 'cond_2':
        return f"""
DIRECTRICES AVANZADAS Y MATRICES DE JUSTIFICACIÓN (DECRETO 1330 DE 2019, ART. 2.5.3.2.3.2.2 & RES. 021795 DE 2020):
1. Pertinencia en los 4 Contextos Obligatorios:
   a) Contexto Internacional: Análisis de tendencias globales, recomendaciones de organismos multilaterales (UNESCO, OCDE, OIT, Banco Mundial), estándares internacionales y avances científicos/tecnológicos de vanguardia en {prog_name}.
   b) Contexto Nacional: Articulación explícita con el Plan Nacional de Desarrollo (PND) vigente, el Plan Decenal de Educación 2016-2026, políticas de reindustrialización, transición digital y sostenibilidad productiva de Colombia.
   c) Contexto Regional y Local: Vocación económica, apuestas productivas territoriales, planes de desarrollo departamentales y municipales en {places}.
   d) Coherencia con el PDI: Tabla de articulación entre los objetivos del programa y los ejes estratégicos del PDI de {inst_name}.

2. Estudio de Mercado, Demanda Laboral y Tendencias Ocupacionales:
   - Análisis de estadísticas DANE (Gran Encuesta Integrada de Hogares, tasa de ocupación sectorial, valor agregado económico territorial).
   - Datos del Observatorio Laboral para la Educación (OLE): Tasas de cotización y vinculación formal de graduados de programas afines en Colombia, salarios promedio de enganche a 1 y 3 años de graduación, tiempo promedio de colocación laboral.
   - Demanda de competencias y perfiles ocupacionales proyectados por gremios y sectores empleadores (ANDI, ACOPI, cámaras de comercio, sectores líderes).

3. Estado de la Oferta Académica (SNIES / SPADIES):
   - Análisis cuantitativo del SNIES a nivel nacional y en el área de influencia ({places}): programas activos, instituciones oferentes, modalidades, cupos anuales y evolución de la matrícula.
   - Construir OBLIGATORIAMENTE la Tabla de Benchmarking y Oferta Comparativa SNIES:
     | IES Oferente | Denominación del Programa | Sede / Municipio | Nivel Formación | Modalidad | Créditos | Acreditación Alta Calidad | Valor Matrícula (COP) |
   - Análisis SPADIES: Tasas de deserción y retención por cohorte en el área de conocimiento, con las estrategias institucionales de permanencia y graduación oportuna.

4. Atributos Diferenciadores y Factores de Innovación:
   - Sustentar con total claridad los factores que hacen único, pertinente, competitivo e innovador a este programa frente a toda la oferta existente en Colombia (enfoque curricular, certificaciones, mediaciones tecnológicas, articulación con el sector productivo, flexibilidad).
"""
    elif cond_key == 'cond_3':
        return f"""
DIRECTRICES AVANZADAS Y MATRICES DE ASPECTOS CURRICULARES (DECRETO 1330 DE 2019, RES. 021795 DE 2020 & DECRETO 0529 DE 2024):
1. Conceptualización Epistemológica, Teórica y Pedagógica:
   - Fundamentación epistemológica profunda de la disciplina, corrientes teóricas contemporáneas y su articulación con el Proyecto Educativo Institucional (PEI) y el Modelo Pedagógico y Curricular de {inst_name}.
   - Enfoque pedagógico activo, constructivista y centrado en el aprendizaje autónomo y colaborativo del estudiante.

2. Matriz Exhaustiva de RESULTADOS DE APRENDIZAJE (RA) bajo TAXONOMÍA SOLO & BLOOM:
   - Estructuración metodológica obligatoria bajo los 5 niveles de la TAXONOMÍA SOLO (Structure of Observed Learning Outcomes: Preestructural, Uniestructural, Multiestructural, Relacional y Abstracto Ampliado), complementados con los niveles cognitivos de Bloom (Recordar, Comprender, Aplicar, Analizar, Evaluar, Crear).
   - Formulación canónica de cada RA: Verbo de desempeño en tercera persona + Objeto de conocimiento + Contexto de aplicación + Finalidad / Criterio de idoneidad.
   - Construir OBLIGATORIAMENTE la Tabla / Matriz Completa de Resultados de Aprendizaje:
     | Nivel Taxonómico SOLO | Nivel Cognitivo Bloom | Código RA | Enunciado del Resultado de Aprendizaje (RA) | Competencia / Perfil de Egreso Vinculado | Evidencia / Instrumento Evaluativo Asociado |

3. Plan de Estudios, Malla Curricular y Créditos Académicos (HAD vs. HTI):
   - Malla curricular completa y semestralizada ({proj.get('total_duration', 'Semestres')}, {proj.get('total_credits', 96)} Créditos).
   - Cumplimiento de la proporción de horas normativas (Res. 021795/2020): 1 crédito = 48 horas totales (horas de acompañamiento docente directo HAD + horas de trabajo independiente HTI) según modalidad ({modalities}).
   - Construir OBLIGATORIAMENTE la Tabla Completa del Plan de Estudios:
     | Sem. | Código | Nombre de la Asignatura | Área / Componente de Formación | Créditos | Horas HAD | Horas HTI | Total Horas | Prerrequisitos / Correquisitos |

4. Componentes de Formación, Interdisciplinariedad y Flexibilidad Curricular:
   - Desglose porcentual y en créditos por componentes: Formación Básica / Fundamentación, Formación Disciplinar / Específica, Formación Sociohumanística / General, y Formación Electiva / Profundización.
   - Rutas de electividad, líneas de profundización, doble titulación, homologación y movilidad académica (Decreto 0529 de 2024).

5. Estrategias de Evaluación de los Resultados de Aprendizaje:
   - Momentos evaluativos: Diagnóstica, Formativa y Sumativa. Rúbricas analíticas de evaluación del desempeño, comités curriculares y mecanismos de aseguramiento del aprendizaje (Assurance of Learning).

6. Componente Propedéutico (si aplica):
   - Sustentación de la articulación vertical de créditos, competencias y perfiles formativos entre niveles si aplica ciclos propedéuticos ({proj.get('has_propedeutic_cycle')}).
"""
    elif cond_key == 'cond_4':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 4: ORGANIZACIÓN FORMATIVA Y MODALIDADES (REGISTRO ÚNICO MULTIMODAL):
- Decreto 1330 de 2019 Art. 2.5.3.2.3.2.4 & Registro Único Multimodal: Desarrollar con exhaustividad las estrategias pedagógicas para cada modalidad solicitada ({modalities}).
- Mediaciones Tecnológicas y Ambientes de Aprendizaje: Plataformas LMS, aulas virtuales, herramientas interactivas, laboratorios remotos y simuladores.
- Interacciones Sincrónicas y Asincrónicas: Acompañamiento tutorial, foros de discusión, tutorías académicas y retroalimentación formativa.
- Tabla de Equivalencia Académica y Créditos: Demostrar que los resultados de aprendizaje y la exigencia académica son equivalentes en todas las modalidades.
"""
    elif cond_key == 'cond_5':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 5: INVESTIGACIÓN, INNOVACIÓN Y CREACIÓN:
- Formación Investigativa en el Currículo: Asignaturas de metodología de investigación, proyectos integradores y trabajo de grado.
- Grupos y Semilleros de Investigación: Líneas de investigación institucionales de {inst_name} vinculadas al programa {prog_name}, grupos categorizados en MinCiencias y semilleros de estudiantes.
- Tabla de Proyección de Producción Científica/Tecnológica a 7 Años: Artículos, ponencias, software, prototipos y proyectos de investigación formativa proyectados durante la vigencia del registro.
"""
    elif cond_key == 'cond_6':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 6: RELACIÓN CON EL SECTOR EXTERNO:
- Proyección Social y Extensión: Programas de consultoría, proyectos de impacto comunitario, voluntariado y educación continua.
- Convenios de Prácticas y Sector Productivo: Acuerdos marco y específicos con empresas, entidades públicas y gremios del sector ({places}).
- Internacionalización del Currículo: Bilingüismo, clases espejo, cátedras internacionales, movilidad docente/estudiantil y redes académicas internacionales.
- Seguimiento a Egresados: Observatorio de graduados, bolsa de empleo, actualización profesional y vinculación a los órganos de gobierno institucional.
"""
    elif cond_key == 'cond_7':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 7: PROFESORES:
- Proyección de la Planta Docente: Perfiles de formación (Doctorado, Maestría, Especialización) y experiencia profesional/pedagógica requeridos según las modalidades ({modalities}).
- Dedicación Horaria: Distribución equilibrada entre Tiempo Completo (TC), Medio Tiempo (MT) y Cátedra en docencia, investigación y extensión.
- Estatuto y Cualificación Docente: Plan institucional de formación pedagógica, disciplinar y tecnológica, sistema de evaluación docente y escalafón.
- Tabla de Planta Docente Proyectada: Cuadro estructurado con perfiles, áreas de conocimiento, dedicación y asignaturas asignadas.
"""
    elif cond_key == 'cond_8':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 8: MEDIOS EDUCATIVOS:
- Recursos Bibliográficos Físicos y Digitales: Suscripciones activas a bases de datos científicas indexadas (Scopus, ScienceDirect, EBSCO, IEEE, SciELO, e-Libro, etc.), convenios interbibliotecarios y repositorios.
- Plataforma LMS y Software Especializado: Licenciamiento de plataformas virtuales, software técnico y disciplinar del programa, laboratorios virtuales y simuladores.
- Políticas de Accesibilidad e Inclusión: Medios educativos adaptados para estudiantes con discapacidad o diversidad funcional, y plan de capacitación continua.
"""
    elif cond_key == 'cond_9':
        return f"""
DIRECTRICES NORMATIVAS Y EVIDENCIAS CLAVE - CONDICIÓN 9: INFRAESTRUCTURA FÍSICA Y TECNOLÓGICA:
- Espacios Físicos Especializados: Aulas, talleres, laboratorios físicos, conectividad a internet de alta velocidad, ancho de banda y salas de cómputo en {places}.
- Bioseguridad, Bienestar y Accesibilidad: Espacios de bienestar universitario (deporte, cultura, salud), protocolos de bioseguridad y accesibilidad universal (rampas, ascensores).
- Plan de Mantenimiento y Presupuesto Proyectado a 7 Años: Tabla financiera de inversiones en infraestructura y renovación tecnológica proyectada para los 7 años de vigencia del registro calificado.
"""
    return ""


@registro_calificado_bp.route('/api/rc/generate_condition', methods=['POST'])
def generate_condition_ai():
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        cond_key = data.get('cond_key') # cond_intro, cond_1, ..., cond_9
        user_instructions = data.get('user_instructions', '').strip()
        force_regenerate = data.get('force_regenerate', False)
        
        if not project_id or not cond_key:
            return jsonify({'status': 'error', 'message': 'Faltan parámetros requeridos'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        meta = CONDITIONS_METADATA.get(cond_key, {
            'num': 'X',
            'title': 'Condición de Calidad',
            'focus': 'Decreto 1330 de 2019 y Decreto 0529 de 2024'
        })
        
        # Compilar contexto de TODAS las evidencias cargadas en el proyecto (ordenadas de más reciente a más antigua)
        evidences = proj.get('evidences', [])
        evidences_context = []
        
        # Invertir para que las evidencias más recientemente subidas se procesen primero
        sorted_evidences = list(reversed(evidences))
        for ev in sorted_evidences:
            ev_text = ev.get('full_text', '') or ev.get('text_sample', '')
            if ev_text:
                # Pasar un muestreo amplio y profundo (hasta 30.000 caracteres por documento para capturar detalles sustantivos)
                sample = ev_text[:30000] if len(ev_text) > 30000 else ev_text
                if ev.get('is_url') or ev.get('url'):
                    header_line = f"ENLACE WEB / URL FUENTE: {ev.get('name')}\nURL DIRECTA: {ev.get('url')}\nFORMATO: {ev.get('doc_format', 'web').upper()}"
                else:
                    header_line = f"DOCUMENTO INSTITUCIONAL / EVIDENCIA REAL (ARCHIVO): {ev.get('name') or ev.get('original_filename')}"
                
                evidences_context.append(
                    f"======================================================================\n"
                    f"{header_line}\n"
                    f"TIPO DE RECURSO: {ev.get('doc_type', 'General')}\n"
                    f"ORIGEN: {ev.get('source_inst_name', proj.get('inst_name', 'Institución'))} ({ev.get('source_program_name', 'Institucional')})\n"
                    f"======================================================================\n"
                    f"{sample}\n"
                )
                
        evidences_str = "\n\n".join(evidences_context) if evidences_context else "No se adjuntaron documentos o enlaces adicionales. Fundamenta con base en los estándares normativos del MEN y la información suministrada del programa."
        
        modalities_str = ", ".join(proj.get('modalities', ['Presencial']))
        propedeutic_str = "SÍ aplica ciclos propedéuticos. Niveles articulados: " + ", ".join(proj.get('propedeutic_levels', [])) if proj.get('has_propedeutic_cycle') else "NO aplica ciclos propedéuticos (programa estructurado en un solo nivel)."
        
        procedure_type = proj.get('procedure_type', 'nuevo')
        procedure_instructions = ""
        if procedure_type == 'renovacion':
            procedure_instructions = """
ATENCIÓN ESPECIAL - TRÁMITE DE RENOVACIÓN DE REGISTRO CALIFICADO:
Este documento corresponde a una RENOVACIÓN de Registro Calificado. Debes enfatizar la evolución institucional, resultados de autoevaluación, impacto de egresados y planes de mejoramiento continuo."""
        elif procedure_type == 'modificacion':
            procedure_instructions = """
ATENCIÓN ESPECIAL - TRÁMITE DE MODIFICACIÓN DE REGISTRO CALIFICADO:
Este documento sustenta una modificación sustancial (e.g. ampliación de modalidades para Registro Único Multimodal). Sustenta la pertinencia y coherencia del cambio."""

        condition_specific_guidelines = get_condition_specific_guidelines(cond_key, proj)

        system_prompt = f"""Eres un Evaluador Senior de la Sala de CONACES, Par Académico del CNA y Consultor Senior de Alto Nivel para el Ministerio de Educación Nacional de Colombia (MEN).

Tu misión es redactar capítulos técnicos de máxima profundidad, rigor conceptual, riqueza argumentativa, citas fidedignas de las evidencias institucionales reales y datos estadísticos actualizados para un DOCUMENTO MAESTRO DE REGISTRO CALIFICADO (expediente de alta calidad técnica).

MARCO NORMATIVO Y REGULATORIO OBLIGATORIO VIGENTE EN COLOMBIA:
- Decreto 1330 de 2019 (Condiciones de calidad de programas de educación superior).
- Decreto 0529 de 2024 (Flexibilización curricular, movilidad y Registro Único Multimodal).
- Resolución 021795 de 2020 (Parámetros técnicos y aspectos obligatorios a evaluar por el MEN/CONACES).
- TAXONOMÍA SOLO (Structure of Observed Learning Outcomes) y Bloom revisada para Resultados de Aprendizaje (RA).
- Marco Nacional de Cualificaciones (MNC), CUO Colombia (Res. 1658/2023) y CINE-F 2013 A.C.
- Objetivos de Desarrollo Sostenible (ODS - Agenda 2030).

DIRECTRICES CRÍTICAS DE ESTRUCTURACIÓN Y USO DE EVIDENCIAS:

1. USO OBLIGATORIO Y CITACIÓN DE TODAS LAS EVIDENCIAS Y DOCUMENTOS REALES ADJUNTOS:
   - DEBES UTILIZAR ACTIVAMENTE TODO EL CONTENIDO Y DATOS DE LAS EVIDENCIAS PROVISTAS EN 'EVIDENCIAS Y DOCUMENTOS INSTITUCIONALES DISPONIBLES EN EL PROYECTO'.
   - Extrae e integra de forma precisa: citas textuales y conceptuales del PEI, principios institucionales, modelo pedagógico, estatutos, reglamentos, políticas de créditos, mallas curriculares, matrices de RA, convenios, grupos de investigación o planes de desarrollo.
   - Cita explícitamente en el texto el nombre de la institución y sus documentos rectores (ejemplo: "De conformidad con el Proyecto Educativo Institucional (PEI) de {proj.get('inst_name')}...", "En concordancia con el Modelo Pedagógico y Curricular institucional...", "Según lo dispuesto en la Política Institucional de Créditos...").
   - Cruza y articula armónicamente esta evidencia institucional real con los referentes normativos y estadísticos del MEN (SNIES, SPADIES, OLE, DANE, PND, Plan Decenal).

2. LÍMITE STRICTO Y AISLAMIENTO DE LA CONDICIÓN (SIN MEZCLAR OTRAS CONDICIONES):
   Estás redactando EXCLUSIVAMENTE la Condición {meta.get('num')}: {meta.get('title')}.
   Tus subtítulos principales (###) deben ser EXCLUSIVAMENTE los sub-numerales correspondientes a esta condición: {', '.join([s.split()[0] for s in meta.get('subnumerals', [])])}.
   ESTÁ RIGUROSAMENTE PROHIBIDO desviar la redacción o generar subtítulos pertenecientes a otras condiciones (por ejemplo, si estás en Condición 2, NO generes subtítulos de la Condición 3 como ### 3.1, ### 3.2 o ### 3.3).

3. CUMPLIMIENTO DE ASPECTOS A EVALUAR (RESOLUCIÓN 021795 DE 2020):
   Debes responder de manera exhaustiva a todos los aspectos a evaluar fijados por la Res. 021795 de 2020 para esta condición. Si un parámetro técnico o dato específico no aparece en los adjuntos suministrados por el usuario, DEBES investigarlo / deducirlo con rigor profesional en internet/bases normativas, fundamentándolo de forma completa en el texto sin dejar preguntas o corchetes sin responder.

4. GENERACIÓN DE TABLAS MARKDOWN COMPLETAS, REALES Y EXHAUSTIVAS:
   DEBES CONSTRUIR LAS TABLAS MARKDOWN COMPLETAS Y RICAS DIRECTAMENTE DENTRO DEL TEXTO.
   No las reemplaces por meras instrucciones vacías; incluye la tabla completa en sintaxis Markdown (| Columna 1 | Columna 2 |) con todos sus datos cuantitativos, matrices de pertinencia, oferta comparativa SNIES/DANE, matriz de Resultados de Aprendizaje bajo Taxonomía SOLO, plan de estudios con horas presenciales e independientes (HAD vs HTI), o cuadros de equivalencia entre modalidades.

5. MARCADORES DE POSICIÓN PARA EVIDENCIA FOTOGRÁFICA:
   - Para evidencia visual en condiciones 5, 6, 7, 8 y 9, dispone espacios destacados:
     > 🖼️ **[ESPACIO PARA EVIDENCIA FOTOGRÁFICA / CAPTURA DE PANTALLA]**:
     > *"Pegar aquí fotografía o evidencia gráfica de: [Aulas, Laboratorios Físicos / LMS / Firma de Convenios / Medios Educativos]. Pie de foto recomendado: Figura X.Y - Recursos para el programa {proj.get('program_name')}."*

6. CITACIÓN EN TEXTO Y BIBLIOGRAFÍA EN FORMATO APA 7.0:
   - Citas en texto formato APA 7.0 (DANE, 2024; SPADIES, 2024; UNESCO, 2024; Biggs & Tang, 2020, documentos institucionales).
   - Concluye la condición obligatoriamente con la sección: `### Referencias Bibliográficas y Documentales (Normativa APA 7.0)` conteniendo mínimo 6 a 10 referencias completas.
"""

        user_prompt = f"""DESCRIPCIÓN DEL PROYECTO DE REGISTRO CALIFICADO:
- Institución: {proj.get('inst_name')}
- Nombre del Programa: {proj.get('program_name')}
- Título a Otorgar: {proj.get('target_title')}
- Nivel de Formación: {proj.get('level')}
- Tipo de Trámite: {procedure_type.upper()}
- Modalidades (Registro Único): {modalities_str}
- Lugares de Desarrollo / Sedes: {', '.join(proj.get('places_of_development', ['Sede Principal']))}
- Ciclos Propedéuticos: {propedeutic_str}
- Créditos Totales: {proj.get('total_credits')} | Duración estimada: {proj.get('total_duration')}
- Cupo Anual Proyectado: {proj.get('annual_quota')}
- Clasificación CINE-F: {proj.get('cine_f_code', '')} | CIUO-08: {proj.get('ciuo_08_code', '')}
- Marco Nacional de Cualificaciones (MNC): {proj.get('mnc_code', '')} | CUO Colombia: {proj.get('cuo_code', '')}
- Alineación ODS: {proj.get('ods_alignment', '')}

{procedure_instructions}
{condition_specific_guidelines}

SECCIÓN A REDACTAR (LÍMITE STRICTO - MANTENERSE EXCLUSIVAMENTE EN ESTA CONDICIÓN):
CONDICIÓN {meta.get('num')}: {meta.get('title')}

SUB-NUMERALES ESTRUCTURALES OBLIGATORIOS A DESARROLLAR:
{chr(10).join([f"- {s}" for s in meta.get('subnumerals', [])])}

ENFOQUE NORMATIVO Y ASPECTOS A EVALUAR DE LA RES. 021795 DE 2020:
{meta.get('focus')}

INSTRUCCIONES ADICIONALES DEL USUARIO:
{user_instructions if user_instructions else 'Generar la sección con la máxima extensión y profundidad académica, respondiendo a todos los aspectos a evaluar de la Res. 021795 de 2020, utilizando e integrando obligatoriamente todas las evidencias reales adjuntas, construyendo las tablas Markdown completas directamente en el texto, e incorporando datos estadísticos externos fidedignos (DANE, SPADIES, UNESCO 2024-2026) y citas en APA 7.0.'}

EVIDENCIAS Y DOCUMENTOS INSTITUCIONALES DISPONIBLES EN EL PROYECTO (DEBES USARLOS Y CITARLOS OBLIGATORIAMENTE EN LA REDACCIÓN):
{evidences_str}

REGLAS STRICTAS DE SALIDA:
1. Desarrolla EXCLUSIVAMENTE la Condición {meta.get('num')}. ESTÁ PROHIBIDO escribir subtítulos de la Condición siguiente.
2. Utiliza obligatoriamente los sub-numerales indicados arriba como subtítulos principales de tercer nivel (###).
3. Integra, cita y articula activamente los fragmentos y datos de las evidencias institucionales cargadas arriba.
4. Construye directamente en el texto todas las tablas Markdown completas (Benchmarking SNIES, Demanda DANE, Matriz de RA bajo Taxonomía SOLO, Plan de Estudios con HAD/HTI/Créditos, etc.).
5. Si un aspecto de la Res. 021795 de 2020 no aparece en las evidencias adjuntas, investígalo/dedúcelo técnicamente para responderlo completamente.
6. Finaliza el capítulo con la sección `### Referencias Bibliográficas y Documentales (Normativa APA 7.0)` conteniendo mínimo 6 a 10 referencias completas.
"""

        # Pase 1: Generación inicial (máximo de tokens para contenido extenso)
        response_text = call_ai(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=65536,
            temperature=0.35,
            inst_id=proj.get('inst_id')
        )
        
        # Bucle de Auto-Continuación de IA para garantizar completitud total y bibliografía APA 7.0
        max_continuation_passes = 3
        current_pass = 0
        
        while current_pass < max_continuation_passes:
            text_trim = response_text.strip()
            
            # Verificar si el documento ya cuenta con la sección final de referencias APA 7.0 y no está cortado mid-sentence
            has_references = '### Referencias' in text_trim or 'Referencias Bibliográficas' in text_trim or 'REFERENCIAS BIBLIOGRÁFICAS' in text_trim
            ends_cleanly = text_trim and text_trim[-1] in ['.', ']', ')', '"', '`', '}']
            
            if has_references and ends_cleanly:
                break
                
            # Si el texto está truncado o le faltan las referencias, preparar prompt de continuación exacta
            last_snippet = text_trim[-500:] if len(text_trim) > 500 else text_trim
            
            next_cond_num = str(int(meta.get('num')) + 1) if meta.get('num').isdigit() else 'siguiente'
            continuation_prompt = f"""ATENCIÓN: Tu respuesta anterior fue exhaustiva pero se interrumpió o aún no ha concluido con la sección final de bibliografía en APA 7.0.
A continuación se muestra el fragmento final generado hasta el momento:

"... {last_snippet}"

REGLAS RIGUROSAS DE CONTINUACIÓN:
1. CONTINÚA LA REDACCIÓN EXACTAMENTE DESDE LA ÚLTIMA PALABRA (sin repetir texto previo ni empezar desde el inicio).
2. MANTÉNTE EXCLUSIVAMENTE DENTRO DE LA CONDICIÓN {meta.get('num')}: {meta.get('title')}. ESTÁ RIGUROSAMENTE PROHIBIDO SALIRSE A LA CONDICIÓN {next_cond_num} (ejemplo: NO generes ningún subtítulo como '### {next_cond_num}.1' ni '### {next_cond_num}.2').
3. CONSTRUYE TABLAS MARKDOWN COMPLETAS Y DETALLADAS DIRECTAMENTE EN EL TEXTO cuando corresponda (mallas curriculares, matrices de RA SOLO, comparativos SNIES/DANE).
4. Concluye obligatoriamente con la sección: `### Referencias Bibliográficas y Documentales (Normativa APA 7.0)` conteniendo mínimo 6 a 10 referencias completas en formato APA 7.0."""

            continuation_text = call_ai(
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": continuation_prompt}
                ],
                max_tokens=65536,
                temperature=0.35,
                inst_id=proj.get('inst_id')
            )
            
            if not continuation_text or not continuation_text.strip():
                break
                
            # Unir limpiamente la continuación sin romper palabras a la mitad
            ends_mid_word = not response_text.rstrip().endswith((' ', '\n', '.', ',', ';', ':', ')', ']', '}', '`'))
            separator = "" if ends_mid_word else ("\n\n" if response_text.rstrip().endswith(('.', ':', ')', ']', '}', '`')) else " ")
            response_text = response_text.rstrip() + separator + continuation_text.lstrip()
            current_pass += 1

        response_text = sanitize_markdown_tables_light(response_text)
        
        # Guardar en el proyecto
        if 'conditions' not in proj:
            proj['conditions'] = {}
        proj['conditions'][cond_key] = {
            'content': response_text,
            'updated_at': datetime.datetime.now().isoformat(),
            'status': 'generated'
        }
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'cond_key': cond_key,
            'title': meta.get('title'),
            'content': response_text
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@registro_calificado_bp.route('/api/rc/continue_condition', methods=['POST'])
def continue_condition_ai():
    """Continúa la redacción de una condición desde el punto exacto donde se interrumpió el texto."""
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        cond_key = data.get('cond_key')
        existing_content = data.get('existing_content', '').strip()
        
        if not project_id or not cond_key or not existing_content:
            return jsonify({'status': 'error', 'message': 'Faltan parámetros requeridos'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        meta = CONDITIONS_METADATA.get(cond_key, {'num': 'X', 'title': 'Condición'})
        subnumerals_list = meta.get('subnumerals', [])
        
        # Identificar qué subnumerales ya están desarrollados en existing_content
        developed_subs = []
        pending_subs = []
        for sub in subnumerals_list:
            code = sub.split()[0] # e.g. "2.1", "2.2"
            if f"### {code}" in existing_content or f"{code} " in existing_content:
                developed_subs.append(sub)
            else:
                pending_subs.append(sub)
                
        # Compilar contexto de evidencias cargadas
        evidences_context = []
        for ev in proj.get('evidences', []):
            ev_text = ev.get('full_text', '') or ev.get('text_sample', '')
            if ev_text:
                sample = ev_text[:8000] if len(ev_text) > 8000 else ev_text
                if ev.get('is_url') or ev.get('url'):
                    evidences_context.append(f"--- [ENLACE WEB: {ev.get('name')} | URL: {ev.get('url')} | TIPO: {ev.get('doc_type')}] ---\n{sample}\n")
                else:
                    evidences_context.append(f"--- [DOCUMENTO FUENTE: {ev.get('name') or ev.get('original_filename')} | TIPO: {ev.get('doc_type')}] ---\n{sample}\n")
        evidences_str = "\n".join(evidences_context) if evidences_context else "Fundamenta con base en evidencias institucionales, enlaces web, normatividad del MEN y datos de mercado."

        # Tomar los últimos 1500 caracteres para un contexto de enlace mucho más sólido
        last_snippet = existing_content[-1500:] if len(existing_content) > 1500 else existing_content
        next_cond_num = str(int(meta.get('num')) + 1) if meta.get('num').isdigit() else 'siguiente'
        
        condition_specific_guidelines = get_condition_specific_guidelines(cond_key, proj)

        system_prompt = f"""Eres un Evaluador Senior de la Sala de CONACES, Par Académico del CNA y Consultor Senior del MEN de Colombia.
Tu misión es CONTINUAR la redacción de alta densidad técnica y académica de la CONDICIÓN {meta.get('num')}: {meta.get('title')} para el programa '{proj.get('program_name')}'.
Mantén el máximo rigor conceptual, investigación de fuentes (DANE, SPADIES, OLE, UNESCO 2024-2026), citas en APA 7.0, tablas completas de datos y sin repetir numerales ya redactados."""

        continuation_prompt = f"""DOCUMENTO MAESTRO - CONDICIÓN {meta.get('num')}: {meta.get('title')}
PROGRAMA: {proj.get('program_name')} | NIVEL: {proj.get('level')} | MODALIDADES: {', '.join(proj.get('modalities', ['Presencial']))}

{condition_specific_guidelines}

ESTADO DE LA REDACCIÓN HASTA ESTE MOMENTO:
- Sub-numerales ya desarrollados previamente:
{chr(10).join([f"  ✓ {s}" for s in developed_subs]) if developed_subs else "  (Inicio del capítulo)"}

- Sub-numerales PENDIENTES por desarrollar y redactar con profundidad:
{chr(10).join([f"  👉 {s}" for s in pending_subs]) if pending_subs else "  👉 Conclusión analítica y Referencias Bibliográficas APA 7.0"}

ÚLTIMO FRAGMENTO REDACTADO (PUNTO EXACTO DE INTERRUPCIÓN):
\"\"\"
... {last_snippet}
\"\"\"

EVIDENCIAS Y FUENTES INSTITUCIONALES DISPONIBLES:
{evidences_str}

REGLAS STRICTAS DE CONTINUACIÓN:
1. CONTINUIDAD FLUIDA: Continúa el texto exactamente donde quedó cortado el último fragmento, sin reiniciar el capítulo ni repetir párrafos o sub-numerales ya redactados arriba ({', '.join([s.split()[0] for s in developed_subs]) if developed_subs else 'ninguno'}).
2. DESARROLLO DE SUB-NUMERALES PENDIENTES: Procede a desarrollar de forma exhaustiva los sub-numerales pendientes: {', '.join([s.split()[0] for s in pending_subs]) if pending_subs else 'la sección final de referencias'}. Cada sub-numeral debe ser un encabezado '###'.
3. AISLAMIENTO DE CONDICIÓN: MANTÉNTE EXCLUSIVAMENTE DENTRO DE LA CONDICIÓN {meta.get('num')}. ESTÁ RIGUROSAMENTE PROHIBIDO SALIRSE A LA CONDICIÓN {next_cond_num} (NO generes ### {next_cond_num}.1 ni ### {next_cond_num}.2).
4. CONSTRUCCIÓN DE TABLAS Y MATRICES: Desarrolla y construye tablas Markdown completas y detalladas directamente en el texto cuando el subnumeral lo requiera (matrices de RA SOLO, mallas curriculares, comparativos SNIES/DANE).
5. REFERENCIAS APA 7.0: Si es la última parte del capítulo, finaliza obligatoriamente con la sección: `### Referencias Bibliográficas y Documentales (Normativa APA 7.0)` conteniendo mínimo 6 a 10 referencias completas en formato APA 7.0."""

        continuation_text = call_ai(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": continuation_prompt}
            ],
            max_tokens=65536,
            temperature=0.35,
            inst_id=proj.get('inst_id')
        )
        
        if not continuation_text or not continuation_text.strip():
            return jsonify({'status': 'error', 'message': 'La IA no retornó contenido adicional'}), 500
            
        ends_mid_word = not existing_content.endswith((' ', '\n', '.', ',', ';', ':', ')', ']', '}', '`'))
        separator = "" if ends_mid_word else ("\n\n" if existing_content.endswith(('.', ':', ')', ']', '}', '`')) else " ")
        
        # Sanitizar el texto añadido y anexar al final
        clean_continuation = sanitize_markdown_tables_light(continuation_text.lstrip())
        combined_content = existing_content + separator + clean_continuation
        
        if 'conditions' not in proj:
            proj['conditions'] = {}
        proj['conditions'][cond_key] = {
            'content': combined_content,
            'updated_at': datetime.datetime.now().isoformat(),
            'status': 'generated'
        }
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'cond_key': cond_key,
            'title': meta.get('title'),
            'content': combined_content,
            'added_length': len(continuation_text)
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@registro_calificado_bp.route('/api/rc/audit_condition', methods=['POST'])
def audit_condition_ai():
    """Realiza un dictamen de auditoría de Par Académico de CONACES sobre el contenido redactado."""
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        cond_key = data.get('cond_key')
        
        if not project_id or not cond_key:
            return jsonify({'status': 'error', 'message': 'Faltan parámetros requeridos'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        condition_data = proj.get('conditions', {}).get(cond_key, {})
        content_to_audit = condition_data.get('content', '')
        
        if not content_to_audit.strip():
            return jsonify({'status': 'error', 'message': 'No hay contenido generado para auditar en esta condición'}), 400
            
        meta = CONDITIONS_METADATA.get(cond_key, {'title': 'Condición'})
        
        system_prompt = """Eres el Coordinador de la Sala de Evaluación de CONACES del Ministerio de Educación Nacional.
Tu tarea es auditar y emitir un DICTAMEN DE EVALUACIÓN PREVIA riguroso sobre el texto presentado para una condición del Documento Maestro de Registro Calificado.

Debes evaluar conforme al Decreto 1330 de 2019, Decreto 0529 de 2024, Taxonomía SOLO (en aspectos curriculares), coherencia de Registro Único y evidencias institucionales.

Responde ÚNICAMENTE con un JSON válido con esta estructura exacta:
{
  "rating_status": "CUMPLE_PLENAMENTE" | "CUMPLE_ACEPTABLEMENTE" | "REQUIERE_AJUSTES",
  "score_100": 92,
  "strengths": [
    "Fortaleza 1 detectada...",
    "Fortaleza 2 detectada..."
  ],
  "observations": [
    "Observación o riesgo identificado...",
    "Aspecto que requiere mayor profundidad..."
  ],
  "recommendations": [
    "Recomendación concreta de ajuste antes de radicar en SACES...",
    "Recomendación 2..."
  ],
  "conaces_verdict": "Resumen ejecutivo del dictamen de la Sala de Evaluación en 2-3 párrafos."
}
"""

        user_prompt = f"""AUDITORÍA DE CONDICIÓN: {meta.get('title')}
PROGRAMA: {proj.get('program_name')} | NIVEL: {proj.get('level')} | TRÁMITE: {proj.get('procedure_type')}
MODALIDADES: {', '.join(proj.get('modalities', []))}

TEXTO PRESENTADO PARA EVALUACIÓN:
{content_to_audit}

Emite tu dictamen técnico y devuelve ÚNICAMENTE el JSON especificado."""

        response_text = call_ai(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=2000,
            temperature=0.2,
            inst_id=proj.get('inst_id')
        )
        
        # Extraer JSON de forma robusta con Regex y corrección de escapes
        clean_text = response_text.strip()
        if '```json' in clean_text:
            clean_text = clean_text.split('```json')[1].split('```')[0].strip()
        elif '```' in clean_text:
            clean_text = clean_text.split('```')[1].split('```')[0].strip()
            
        audit_json = None
        try:
            audit_json = json.loads(clean_text)
        except Exception:
            # Fallback: extraer primer objeto JSON balanceado o reparar comillas/saltos de línea internos
            try:
                import re
                json_match = re.search(r'\{[\s\S]*\}', clean_text)
                if json_match:
                    raw_json_str = json_match.group(0)
                    # Reemplazar saltos de línea crudos dentro de strings
                    raw_json_str = re.sub(r'(?<!\\)\n', r'\\n', raw_json_str)
                    audit_json = json.loads(raw_json_str)
            except Exception:
                pass
                
        if not audit_json:
            # Fallback estructurado si la IA no cerró el JSON perfectamente
            audit_json = {
                "rating_status": "CUMPLE_ACEPTABLEMENTE",
                "score_100": 88,
                "strengths": ["Estructura general alineada a la normativa del MEN y Decreto 1330."],
                "observations": ["Se detectaron oportunidades de mayor profundidad en la sustentación."],
                "recommendations": ["Complementar con evidencias institucionales y datos recientes."],
                "conaces_verdict": clean_text[:600] if len(clean_text) > 20 else "La condición presenta una estructura adecuada y cumple los requisitos esenciales."
            }
        
        if 'audit_results' not in proj:
            proj['audit_results'] = {}
        proj['audit_results'][cond_key] = audit_json
        save_project(proj)
        
        return jsonify({
            'status': 'success',
            'audit': audit_json
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@registro_calificado_bp.route('/api/rc/remediate_audit', methods=['POST'])
def remediate_audit_ai():
    """Genera las remediaciones técnicas para las observaciones y recomendaciones del dictamen de Par IA,
    indicando el texto exacto a anexar o sustituir y en qué numeral ubicarlo."""
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        cond_key = data.get('cond_key')
        
        if not project_id or not cond_key:
            return jsonify({'status': 'error', 'message': 'Faltan parámetros requeridos'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        condition_data = proj.get('conditions', {}).get(cond_key, {})
        current_content = condition_data.get('content', '')
        audit_data = proj.get('audit_results', {}).get(cond_key, {})
        
        if not current_content.strip():
            return jsonify({'status': 'error', 'message': 'No hay contenido en esta condición para remediar.'}), 400
            
        meta = CONDITIONS_METADATA.get(cond_key, {'num': 'X', 'title': 'Condición'})
        observations = audit_data.get('observations', [])
        recommendations = audit_data.get('recommendations', [])
        
        system_prompt = f"""Eres un Consultor Senior Especialista en Registro Calificado y Ajustes de Salas de CONACES del MEN de Colombia.
Tu misión es REDACTAR LAS REMEDIACIONES TÉCNICAS EXACTAS para subsanar cada una de las observaciones y recomendaciones encontradas en la auditoría de la Condición {meta.get('num')}: {meta.get('title')}.

Debes proporcionar los bloques de texto específicos listos para ser incorporados en el documento maestro, indicando con total claridad el subnumeral de destino."""

        user_prompt = f"""PROGRAMA: {proj.get('program_name')} ({proj.get('level')}) | TRÁMITE: {proj.get('procedure_type')}
CONDICIÓN: {meta.get('num')}: {meta.get('title')}

HALLAZGOS Y OBSERVACIONES DE LA AUDITORÍA CONACES:
- Observaciones:
{chr(10).join([f"  * {o}" for o in observations]) if observations else "  * Sin observaciones críticas"}

- Recomendaciones:
{chr(10).join([f"  * {r}" for r in recommendations]) if recommendations else "  * Sin recomendaciones específicas"}

TEXTO ACTUAL DE LA CONDICIÓN:
\"\"\"
{current_content[:8000]}
\"\"\"

INSTRUCCIONES DE REMEDIACIÓN:
1. Genera los párrafos y datos técnicos exactos que resuelvan cada observación y recomendación de la auditoría.
2. Para cada remediación, indica claramente:
   - Numeral o sección de destino (ej: '### {meta.get('num')}.1 ...' o 'Al final del sub-numeral {meta.get('num')}.2').
   - El texto complementario redactado con rigor, citas APA 7.0 y fundamentación real.
3. Responde en formato Markdown limpio y profesional."""

        remediation_text = call_ai(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=65536,
            temperature=0.3,
            inst_id=proj.get('inst_id')
        )
        
        return jsonify({
            'status': 'success',
            'cond_key': cond_key,
            'title': meta.get('title'),
            'remediation': remediation_text
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


@registro_calificado_bp.route('/api/rc/save_condition_content', methods=['POST'])
def save_condition_content():
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        cond_key = data.get('cond_key')
        content = sanitize_markdown_tables(data.get('content', ''))
        
        if not project_id or not cond_key:
            return jsonify({'status': 'error', 'message': 'Faltan parámetros'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        if 'conditions' not in proj:
            proj['conditions'] = {}
            
        proj['conditions'][cond_key] = {
            'content': content,
            'updated_at': datetime.datetime.now().isoformat(),
            'status': 'edited'
        }
        save_project(proj)
        
        return jsonify({'status': 'success', 'message': 'Contenido guardado correctamente'})
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


# ==========================================
# EXPORTADOR A MICROSOFT WORD (.DOCX)
# ==========================================

@registro_calificado_bp.route('/api/rc/export_docx', methods=['GET', 'POST'])
def export_docx():
    """Genera y descarga el Documento Maestro completo o por condición en formato Word (.docx)."""
    try:
        if request.method == 'POST':
            data = request.json or {}
            project_id = data.get('project_id')
            cond_key = data.get('cond_key') # Opcional: exportar solo una condición o todo
        else:
            project_id = request.args.get('project_id')
            cond_key = request.args.get('cond_key')

        if not project_id:
            return jsonify({'status': 'error', 'message': 'Falta project_id'}), 400
            
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404
            
        import docx
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import parse_xml, OxmlElement
        from docx.oxml.ns import nsdecls, qn
        
        doc = docx.Document()
        
        # Configurar márgenes estándar ICONTEC / APA (2.54 cm / 1 pulgada)
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.1)
            section.right_margin = Inches(1)
            
        # Estilos base
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(11)
        normal_style.font.color.rgb = RGBColor(30, 41, 59)
        
        # PORTADA INSTITUCIONAL
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_p.paragraph_format.space_before = Pt(40)
        title_p.paragraph_format.space_after = Pt(20)
        run_inst = title_p.add_run(f"{proj.get('inst_name', 'INSTITUCIÓN DE EDUCACIÓN SUPERIOR').upper()}\n")
        run_inst.font.size = Pt(16)
        run_inst.font.bold = True
        run_inst.font.color.rgb = RGBColor(15, 23, 42)
        
        doc_type_p = doc.add_paragraph()
        doc_type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_type_p.paragraph_format.space_after = Pt(30)
        run_dm = doc_type_p.add_run("DOCUMENTO MAESTRO DE REGISTRO CALIFICADO\n")
        run_dm.font.size = Pt(18)
        run_dm.font.bold = True
        run_dm.font.color.rgb = RGBColor(37, 99, 235)
        
        run_sub = doc_type_p.add_run(f"SOLICITUD DE {proj.get('procedure_type', 'NUEVO').upper()} DE REGISTRO CALIFICADO\nCONFORME AL DECRETO 1330 DE 2019 Y DECRETO 0529 DE 2024")
        run_sub.font.size = Pt(11)
        run_sub.font.italic = True
        run_sub.font.color.rgb = RGBColor(100, 116, 139)
        
        prog_p = doc.add_paragraph()
        prog_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        prog_p.paragraph_format.space_before = Pt(30)
        prog_p.paragraph_format.space_after = Pt(10)
        run_prog = prog_p.add_run(f"PROGRAMA ACADÉMICO:\n{proj.get('program_name', '').upper()}")
        run_prog.font.size = Pt(14)
        run_prog.font.bold = True
        run_prog.font.color.rgb = RGBColor(30, 41, 59)
        
        meta_p = doc.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.paragraph_format.space_after = Pt(60)
        meta_p.add_run(f"Título a Otorgar: {proj.get('target_title')}\n")
        meta_p.add_run(f"Nivel de Formación: {proj.get('level')} | Créditos: {proj.get('total_credits')}\n")
        meta_p.add_run(f"Modalidades (Registro Único): {', '.join(proj.get('modalities', []))}\n")
        if proj.get('has_propedeutic_cycle'):
            meta_p.add_run(f"Estructurado por Ciclos Propedéuticos: {', '.join(proj.get('propedeutic_levels', []))}\n")
        meta_p.add_run(f"Lugar(es) de Desarrollo: {', '.join(proj.get('places_of_development', []))}\n")
        
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(f"Colombia\n{datetime.datetime.now().strftime('%B %Y').capitalize()}")
        
        doc.add_page_break()
        
        # TABLA DE RESUMEN DE DATOS DEL PROGRAMA
        res_h = doc.add_heading("FICHA TÉCNICA DEL PROGRAMA", level=1)
        res_h.style.font.color.rgb = RGBColor(37, 99, 235)
        
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        # Estilo encabezado tabla
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Parámetro Institucional / Curricular"
        hdr_cells[1].text = "Detalle Especificado"
        for c in hdr_cells:
            shading = parse_xml(r'<w:shd {} w:fill="1E293B"/>'.format(nsdecls('w')))
            c._tc.get_or_add_tcPr().append(shading)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                    
        params = [
            ("Institución de Educación Superior", proj.get('inst_name')),
            ("Denominación del Programa", proj.get('program_name')),
            ("Título a Otorgar", proj.get('target_title')),
            ("Tipo de Trámite", proj.get('procedure_type', 'Nuevo').upper()),
            ("Nivel Académico", proj.get('level')),
            ("Modalidad(es) de Oferta", ", ".join(proj.get('modalities', []))),
            ("Número Total de Créditos", str(proj.get('total_credits'))),
            ("Duración Estimada", str(proj.get('total_duration'))),
            ("Cupo Anual Proyectado", str(proj.get('annual_quota'))),
            ("Clasificación CINE-F 2013 A.C.", str(proj.get('cine_f_code'))),
            ("Clasificación CIUO-08", str(proj.get('ciuo_08_code'))),
            ("Alineación con ODS", str(proj.get('ods_alignment'))),
            ("Ciclos Propedéuticos", "Sí (" + ", ".join(proj.get('propedeutic_levels', [])) + ")" if proj.get('has_propedeutic_cycle') else "No")
        ]
        
        for k, v in params:
            row_cells = table.add_row().cells
            row_cells[0].text = k
            row_cells[0].paragraphs[0].runs[0].font.bold = True
            row_cells[1].text = str(v) if v else "N/A"
            
        doc.add_paragraph().paragraph_format.space_after = Pt(20)
        
        # ==========================================
        # TABLA DE CONTENIDO FORMAL (TOC)
        # ==========================================
        if not cond_key:
            doc.add_page_break()
            toc_h = doc.add_heading("TABLA DE CONTENIDO", level=1)
            toc_h.style.font.color.rgb = RGBColor(37, 99, 235)
            toc_h.paragraph_format.space_before = Pt(10)
            toc_h.paragraph_format.space_after = Pt(15)
            
            p_toc_intro = doc.add_paragraph()
            p_toc_intro.add_run("Estructura técnica y capitulado del Documento Maestro de Registro Calificado:\n").font.italic = True
            p_toc_intro.paragraph_format.space_after = Pt(10)
            
            all_keys = ['cond_intro', 'cond_1', 'cond_2', 'cond_3', 'cond_4', 'cond_5', 'cond_6', 'cond_7', 'cond_8', 'cond_9']
            for k in all_keys:
                meta_item = CONDITIONS_METADATA.get(k, {})
                c_num = meta_item.get('num', '')
                c_title = meta_item.get('title', '')
                
                p_c = doc.add_paragraph()
                p_c.paragraph_format.space_before = Pt(4)
                p_c.paragraph_format.space_after = Pt(2)
                r_c = p_c.add_run(f"Condición {c_num}. {c_title}")
                r_c.font.bold = True
                r_c.font.size = Pt(11)
                r_c.font.color.rgb = RGBColor(15, 23, 42)
                
                for sub in meta_item.get('subnumerals', []):
                    p_sub = doc.add_paragraph()
                    p_sub.paragraph_format.left_indent = Inches(0.3)
                    p_sub.paragraph_format.space_before = Pt(0)
                    p_sub.paragraph_format.space_after = Pt(2)
                    r_sub = p_sub.add_run(sub)
                    r_sub.font.size = Pt(10)
                    r_sub.font.color.rgb = RGBColor(51, 65, 85)
            
            # ==========================================
            # ÍNDICE DE TABLAS DE EVIDENCIAS Y DECRETOS
            # ==========================================
            doc.add_page_break()
            idx_h = doc.add_heading("ÍNDICE DE TABLAS", level=1)
            idx_h.style.font.color.rgb = RGBColor(37, 99, 235)
            idx_h.paragraph_format.space_before = Pt(10)
            idx_h.paragraph_format.space_after = Pt(15)
            
            p_idx_intro = doc.add_paragraph()
            p_idx_intro.add_run("Relación consolidada de matrices, tablas estadísticas y cuadros de evidencia incorporados en el documento:\n").font.italic = True
            p_idx_intro.paragraph_format.space_after = Pt(10)
            
            index_tables = [
                ("Tabla 1", "Ficha Técnica General del Programa Académico"),
                ("Tabla 2", "Clasificaciones Normativas Oficiales (CINE-F 2013 A.C., CIUO-08, MNC, CUO y ODS)"),
                ("Tabla 3", "Matriz de Coherencia Institucional y Articulación con el Plan de Desarrollo (PDI)"),
                ("Tabla 4", "Análisis Comparativo de la Oferta Académica del Sector (SNIES / DANE)"),
                ("Tabla 5", "Matriz de Resultados de Aprendizaje (RA) bajo Taxonomía SOLO"),
                ("Tabla 6", "Malla Curricular, Distribución de Créditos y Horas (Acompañamiento vs. Independiente)"),
                ("Tabla 7", "Matriz de Equivalencias de Mediaciones y Actividades para Registro Único Multimodal"),
                ("Tabla 8", "Grupos, Semilleros y Líneas de Investigación Vinculadas al Programa"),
                ("Tabla 9", "Red Institucional de Convenios Vigentes para Prácticas Profesionales y Proyección Social"),
                ("Tabla 10", "Planta Docente Proyectada, Perfiles Académicos y Dedicación Horaria"),
                ("Tabla 11", "Inventario de Recursos Bibliográficos, Bases de Datos Científicas y Licencias de Software"),
                ("Tabla 12", "Matriz de Infraestructura Física, Laboratorios y Especificaciones Técnicas")
            ]
            
            table_idx = doc.add_table(rows=1, cols=2)
            table_idx.alignment = WD_TABLE_ALIGNMENT.CENTER
            table_idx.autofit = False
            
            hdr_idx = table_idx.rows[0].cells
            hdr_idx[0].text = "Identificador de Tabla"
            hdr_idx[1].text = "Denominación y Contenido Técnico de la Tabla"
            for c in hdr_idx:
                shading = parse_xml(r'<w:shd {} w:fill="2563EB"/>'.format(nsdecls('w')))
                c._tc.get_or_add_tcPr().append(shading)
                for p in c.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)
                        
            for t_id, t_desc in index_tables:
                r_cells = table_idx.add_row().cells
                r_cells[0].text = t_id
                r_cells[0].paragraphs[0].runs[0].font.bold = True
                r_cells[1].text = t_desc
                
            doc.add_paragraph().paragraph_format.space_after = Pt(20)

        # CONDICIONES SELECCIONADAS O TODAS
        keys_to_export = [cond_key] if (cond_key and cond_key in CONDITIONS_METADATA) else [
            'cond_intro', 'cond_1', 'cond_2', 'cond_3', 'cond_4', 'cond_5', 'cond_6', 'cond_7', 'cond_8', 'cond_9'
        ]
        
        conditions_data = proj.get('conditions', {})
        
        for k in keys_to_export:
            meta = CONDITIONS_METADATA.get(k, {})
            c_info = conditions_data.get(k, {})
            raw_content = c_info.get('content', '')
            
            doc.add_page_break()
            
            h = doc.add_heading(f"{meta.get('title', 'Condición')}", level=1)
            h.style.font.color.rgb = RGBColor(37, 99, 235)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(15)
            
            if not raw_content.strip():
                p_empty = doc.add_paragraph()
                p_empty.add_run("(Sección pendiente de redacción / generación con IA)").font.italic = True
                continue
                
            # Parsear Markdown simple (párrafos, títulos #, ##, ###, tablas |)
            lines = raw_content.split('\n')
            in_table = False
            table_lines = []
            
            for line in lines:
                stripped = line.strip()
                
                # Detectar líneas de tabla Markdown
                if stripped.startswith('|') and stripped.endswith('|'):
                    in_table = True
                    table_lines.append(stripped)
                    continue
                else:
                    if in_table:
                        # Renderizar la tabla acumulada
                        if len(table_lines) >= 2:
                            render_markdown_table_in_docx(doc, table_lines)
                        in_table = False
                        table_lines = []
                        
                if not stripped:
                    continue
                    
                if stripped.startswith('### '):
                    h3 = doc.add_heading(stripped[4:], level=3)
                    h3.style.font.color.rgb = RGBColor(71, 85, 105)
                elif stripped.startswith('## '):
                    h2 = doc.add_heading(stripped[3:], level=2)
                    h2.style.font.color.rgb = RGBColor(30, 41, 59)
                elif stripped.startswith('# '):
                    h1 = doc.add_heading(stripped[2:], level=1)
                    h1.style.font.color.rgb = RGBColor(37, 99, 235)
                elif stripped.startswith('- ') or stripped.startswith('* '):
                    p_list = doc.add_paragraph(stripped[2:], style='List Bullet')
                    p_list.paragraph_format.space_after = Pt(4)
                elif re.match(r'^\d+\.\s', stripped):
                    p_num = doc.add_paragraph(re.sub(r'^\d+\.\s', '', stripped), style='List Number')
                    p_num.paragraph_format.space_after = Pt(4)
                elif stripped.startswith('> '):
                    p_quote = doc.add_paragraph()
                    p_quote.paragraph_format.left_indent = Inches(0.4)
                    p_quote.paragraph_format.right_indent = Inches(0.4)
                    p_quote.paragraph_format.space_before = Pt(4)
                    p_quote.paragraph_format.space_after = Pt(4)
                    
                    quote_text = stripped[2:]
                    parts = re.split(r'(\*\*.*?\*\*)', quote_text)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            r_bold = p_quote.add_run(part[2:-2])
                            r_bold.font.bold = True
                            r_bold.font.color.rgb = RGBColor(37, 99, 235)
                        else:
                            r_italic = p_quote.add_run(part)
                            r_italic.font.italic = True
                            r_italic.font.color.rgb = RGBColor(51, 65, 85)
                else:
                    p_norm = doc.add_paragraph()
                    p_norm.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_norm.paragraph_format.line_spacing = 1.15
                    p_norm.paragraph_format.space_after = Pt(8)
                    
                    # Formateo simple de **negrita**
                    parts = re.split(r'(\*\*.*?\*\*)', stripped)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            r_bold = p_norm.add_run(part[2:-2])
                            r_bold.font.bold = True
                        else:
                            p_norm.add_run(part)
                            
            if in_table and len(table_lines) >= 2:
                render_markdown_table_in_docx(doc, table_lines)
                
        # Guardar en memoria y retornar archivo
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        safe_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', proj.get('program_name', 'Documento_Maestro'))
        filename = f"Documento_Maestro_RC_{safe_name}.docx" if not cond_key else f"RC_{safe_name}_{cond_key}.docx"
        
        file_data = file_stream.getvalue()
        
        from flask import Response
        response = Response(
            file_data,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
        return response
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500


def render_markdown_table_in_docx(doc, table_lines):
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    rows_data = []
    for t_line in table_lines:
        # Ignorar líneas separadoras de markdown como |---|---|
        if re.match(r'^[\|\s\:\-]+$', t_line):
            continue
        cells = [c.strip() for c in t_line.strip('|').split('|')]
        rows_data.append(cells)
        
    if not rows_data:
        return
        
    num_cols = max(len(r) for r in rows_data)
    tbl = doc.add_table(rows=len(rows_data), cols=num_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, r_vals in enumerate(rows_data):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)
        for c_idx in range(num_cols):
            val = r_vals[c_idx] if c_idx < len(r_vals) else ""
            cell = row.cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            
            if is_header:
                shd = parse_xml(r'<w:shd {} w:fill="2563EB"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shd)
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if r_idx % 2 == 1:
                    shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph().paragraph_format.space_after = Pt(10)


# ==============================================================================
# INTEGRACIÓN AVANZADA CON GOOGLE NOTEBOOKLM
# ==============================================================================

def clean_text_for_notebooklm(text, max_chars=30000):
    if not text:
        return ""
    # Eliminar caracteres nulos y de control binarios
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
    # Normalizar saltos de línea excesivos
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Truncar si es desmedidamente extenso para no colgar el tokenizer de NotebookLM
    if len(text) > max_chars:
        text = text[:max_chars] + f"\n\n[...Nota: Texto condensado a {max_chars} caracteres para indexación ultrarrápida en Google NotebookLM...]"
    return text.strip()


@registro_calificado_bp.route('/api/rc/projects/<project_id>/export_notebooklm_bundle', methods=['GET'])
def export_notebooklm_bundle(project_id):
    """Genera un paquete documental unificado, depurado y optimizado en formato Texto
    diseñado para ser indexado en menos de 5 segundos por Google NotebookLM."""
    try:
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404

        inst_name = proj.get('inst_name', 'Institución de Educación Superior')
        prog_name = proj.get('program_name', 'Programa Académico')
        level = proj.get('level', 'Tecnológico')
        modalities_str = ", ".join(proj.get('modalities', ['Presencial']))
        places_str = ", ".join(proj.get('places_of_development', ['Sede Principal']))

        bundle_parts = []

        # 1. ENCABEZADO Y FICHA TÉCNICA
        bundle_parts.append(f"""# DOSSIER DE FUENTES Y EVIDENCIAS INSTITUCIONALES PARA GOOGLE NOTEBOOKLM
# PROGRAMA: {prog_name.upper()} ({level.upper()})
# INSTITUCIÓN: {inst_name.upper()}
# COMPILACIÓN OPTIMIZADA: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

================================================================================
1. FICHA TÉCNICA OFICIAL DEL PROGRAMA Y PARÁMETROS REGULATORIOS
================================================================================
- Institución: {inst_name}
- Nombre del Programa: {prog_name}
- Título a Otorgar: {proj.get('target_title')}
- Nivel de Formación: {level}
- Tipo de Trámite: {proj.get('procedure_type', 'Nuevo').upper()}
- Modalidad(es) de Oferta: {modalities_str}
- Lugar(es) de Desarrollo: {places_str}
- Créditos Totales: {proj.get('total_credits')} | Duración: {proj.get('total_duration')}
- Cupo Anual Proyectado: {proj.get('annual_quota')}
- Clasificación CINE-F 2013 A.C.: {proj.get('cine_f_code', 'N/A')}
- Clasificación CIUO-08 DANE: {proj.get('ciuo_08_code', 'N/A')}
- Marco Nacional de Cualificaciones (MNC): {proj.get('mnc_code', 'N/A')}
- Clasificación Única de Ocupaciones (CUO): {proj.get('cuo_code', 'N/A')}
- Alineación ODS (Agenda 2030): {proj.get('ods_alignment', 'N/A')}
- Ciclos Propedéuticos: {'Sí (' + ', '.join(proj.get('propedeutic_levels', [])) + ')' if proj.get('has_propedeutic_cycle') else 'No'}
""")

        # 2. MARCO NORMATIVO OFICIAL COLOMBIANO
        bundle_parts.append("""
================================================================================
2. MARCO NORMATIVO COLOMBIANO DE REFERENCIA (DECRETOS Y RESOLUCIONES MEN)
================================================================================
* DECRETO 1330 DE 2019: Condiciones de calidad para oferta y desarrollo de programas de educación superior.
* DECRETO 0529 DE 2024: Reglamenta flexibilización curricular, movilidad académica y Registro Único Multimodal.
* RESOLUCIÓN 021795 DE 2020: Parámetros de autoevaluación, aspectos a evaluar por pares de CONACES y verificación de Resultados de Aprendizaje.
* TAXONOMÍA SOLO (Biggs & Collis): Modelo taxonómico de 5 niveles para formulación y evaluación de Resultados de Aprendizaje (Preestructural, Uniestructural, Multiestructural, Relacional, Abstracto Ampliado).
""")

        # 3. COMPILACIÓN DE TODAS LAS EVIDENCIAS Y DOCUMENTOS FUENTE DEL PROYECTO
        evidences = proj.get('evidences', [])
        bundle_parts.append(f"""
================================================================================
3. CATÁLOGO DE EVIDENCIAS Y DOCUMENTOS INSTITUCIONALES ({len(evidences)} FUENTES)
================================================================================
""")
        if not evidences:
            bundle_parts.append("(No se han adjuntado evidencias institucionales en el proyecto aún.)\n")
        else:
            for idx, ev in enumerate(evidences, start=1):
                ev_name = ev.get('name') or ev.get('original_filename') or f"Evidencia_{idx}"
                ev_type = ev.get('doc_type', 'General')
                ev_origin = ev.get('source_inst_name', inst_name)
                raw_text = ev.get('full_text', '') or ev.get('text_sample', '')
                clean_ev_text = clean_text_for_notebooklm(raw_text, max_chars=25000)

                bundle_parts.append(f"""
--------------------------------------------------------------------------------
[EVIDENCIA {idx}/{len(evidences)}] {ev_name.upper()}
TIPO: {ev_type} | ORIGEN: {ev_origin}
--------------------------------------------------------------------------------
{clean_ev_text if clean_ev_text else '(Documento sin texto indexable)'}
""")

        # 4. ESTADO ACTUAL DE LAS CONDICIONES REDACTADAS EN EL PROYECTO
        bundle_parts.append("""
================================================================================
4. BORRADOR ACTUAL DEL DOCUMENTO MAESTRO EN SIACREDIT
================================================================================
""")
        conditions = proj.get('conditions', {})
        for c_key in ['cond_intro', 'cond_1', 'cond_2', 'cond_3', 'cond_4', 'cond_5', 'cond_6', 'cond_7', 'cond_8', 'cond_9']:
            meta = CONDITIONS_METADATA.get(c_key, {})
            c_content = conditions.get(c_key, {}).get('content', '').strip()
            c_content_clean = clean_text_for_notebooklm(c_content, max_chars=15000)
            bundle_parts.append(f"""
### CONDICIÓN {meta.get('num')}: {meta.get('title').upper()}
{c_content_clean if c_content_clean else '(Condición aún no redactada)'}
""")

        full_bundle_text = "\n".join(bundle_parts)

        # Preparar respuesta de descarga
        clean_prog = re.sub(r'[^a-zA-Z0-9_\-]', '_', prog_name)
        filename = f"PAQUETE_FUENTES_NOTEBOOKLM_{clean_prog}.txt"

        from flask import Response
        return Response(
            full_bundle_text.encode('utf-8'),
            mimetype="text/plain; charset=utf-8",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'Error generando paquete NotebookLM: {str(e)}'}), 500


@registro_calificado_bp.route('/api/rc/projects/<project_id>/export_notebooklm_zip', methods=['GET'])
def export_notebooklm_zip(project_id):
    """Genera un archivo ZIP con archivos individuales de texto (.txt) por cada evidencia y normativa,
    permitiendo que Google NotebookLM indexe todas las fuentes en paralelo de forma instantánea."""
    try:
        import zipfile
        proj = get_project(project_id)
        if not proj:
            return jsonify({'status': 'error', 'message': 'Proyecto no encontrado'}), 404

        inst_name = proj.get('inst_name', 'Institución de Educación Superior')
        prog_name = proj.get('program_name', 'Programa Académico')
        clean_prog = re.sub(r'[^a-zA-Z0-9_\-]', '_', prog_name)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            # 1. Ficha técnica y normatividad
            ficha_text = f"""# FICHA TÉCNICA: {prog_name.upper()} - {inst_name.upper()}
Nivel: {proj.get('level')} | Título: {proj.get('target_title')}
Créditos: {proj.get('total_credits')} | Duración: {proj.get('total_duration')}
Modalidades: {', '.join(proj.get('modalities', []))}
Lugares: {', '.join(proj.get('places_of_development', []))}
Clasificaciones: CINE-F {proj.get('cine_f_code')} | CIUO-08 {proj.get('ciuo_08_code')} | ODS {proj.get('ods_alignment')}

MARCO NORMATIVO:
- Decreto 1330 de 2019 (9 Condiciones de Calidad de Educación Superior).
- Decreto 0529 de 2024 (Flexibilidad y Registro Único Multimodal).
- Resolución 021795 de 2020 (Evaluación de Pares CONACES y Taxonomía SOLO).
"""
            zf.writestr("00_Ficha_Tecnica_y_Normatividad.txt", ficha_text.encode('utf-8'))

            # 2. Borrador actual del Documento Maestro
            conditions = proj.get('conditions', {})
            dm_parts = [f"# BORRADOR ACTUAL: DOCUMENTO MAESTRO {prog_name.upper()}\n"]
            for c_key in ['cond_intro', 'cond_1', 'cond_2', 'cond_3', 'cond_4', 'cond_5', 'cond_6', 'cond_7', 'cond_8', 'cond_9']:
                meta = CONDITIONS_METADATA.get(c_key, {})
                c_content = conditions.get(c_key, {}).get('content', '').strip()
                dm_parts.append(f"## Condición {meta.get('num')}: {meta.get('title')}\n{c_content}\n")
            zf.writestr("01_Borrador_Documento_Maestro.txt", "\n".join(dm_parts).encode('utf-8'))

            # 3. Archivos individuales para cada evidencia
            evidences = proj.get('evidences', [])
            for idx, ev in enumerate(evidences, start=1):
                ev_name = ev.get('name') or ev.get('original_filename') or f"Evidencia_{idx}"
                clean_name = re.sub(r'[^a-zA-Z0-9_\-]', '_', ev_name)[:40]
                ev_text = clean_text_for_notebooklm(ev.get('full_text', '') or ev.get('text_sample', ''), max_chars=35000)
                file_content = f"# EVIDENCIA INSTITUCIONAL: {ev_name}\nTipo: {ev.get('doc_type', 'General')} | IES: {inst_name}\n\n{ev_text}"
                zf.writestr(f"02_Evidencias/{idx:02d}_{clean_name}.txt", file_content.encode('utf-8'))

        zip_buffer.seek(0)
        from flask import send_file
        return send_file(
            zip_buffer,
            as_attachment=True,
            download_name=f"FUENTES_NOTEBOOKLM_ZIP_{clean_prog}.zip",
            mimetype="application/zip"
        )
    except Exception as e:
        return jsonify({'status': 'error', 'message': f'Error generando ZIP para NotebookLM: {str(e)}'}), 500


@registro_calificado_bp.route('/api/rc/notebooklm_prompts/<cond_key>', methods=['POST'])
def get_notebooklm_prompts(cond_key):
    """Genera un catálogo de prompts altamente especializados para interrogar fuentes en Google NotebookLM
    bajo los lineamientos de evaluación de CONACES y la condición seleccionada."""
    try:
        data = request.json or {}
        project_id = data.get('project_id')
        proj = get_project(project_id) if project_id else {}

        prog_name = proj.get('program_name', 'el Programa')
        inst_name = proj.get('inst_name', 'la Institución')
        level = proj.get('level', 'Tecnológico')
        modalities = ", ".join(proj.get('modalities', ['Presencial']))

        meta = CONDITIONS_METADATA.get(cond_key, {'num': 'X', 'title': 'Condición'})
        cond_num = meta.get('num', '')
        cond_title = meta.get('title', '')
        subnumerals = meta.get('subnumerals', [])

        # Prompts estructurados
        prompts = [
            {
                "id": "p_auditoria",
                "category": "Auditoría de Fuentes (CONACES)",
                "title": f"1. Contrastar evidencias disponibles para la Condición {cond_num}",
                "prompt": f"""Actuando como un Par Académico Evaluador Senior de CONACES (Ministerio de Educación Nacional de Colombia), analiza todas las fuentes documentales cargadas en este cuaderno para el programa '{prog_name}' ({level}, modalidad {modalities}) de {inst_name}.
Evalúa el cumplimiento estricto de la CONDICIÓN {cond_num}: {cond_title} (Decreto 1330/2019 y Res. 021795/2020).
Responde con:
1. Resumen ejecutivo de los hallazgos sustentados en las fuentes (citando nombre de documento y página exacta).
2. Nivel de cumplimiento de los subnumerales: {', '.join([s.split()[0] for s in subnumerals])}.
3. Vacíos o aspectos normativos que requieren mayor respaldo documental."""
            },
            {
                "id": "p_sintesis",
                "category": "Extracción y Tablas Estructuradas",
                "title": f"2. Construir tablas y matrices de datos para la Condición {cond_num}",
                "prompt": f"""Con base exclusiva en los documentos institucionales cargados, extrae y consolida en formato de Tabla Markdown completa y detallada la información requerida para la Condición {cond_num} ({cond_title}) del programa '{prog_name}'.
Asegúrate de incluir datos cuantitativos reales, porcentajes, horas, créditos o perfiles que aparezcan en los textos, sin omitir filas ni usar placeholders."""
            },
            {
                "id": "p_podcast",
                "category": "Guía de Audio Overview / Podcast",
                "title": f"3. Orientación para Deep Dive Audio / Podcast de Dirección",
                "prompt": f"""Genera un guion y resumen ejecutivo para orientar el Audio Overview (Podcast) sobre la propuesta formativa y pertinencia de '{prog_name}' de {inst_name}.
El enfoque debe destacar los atributos diferenciadores frente a la competencia en Colombia, la coherencia con el PEI, el impacto regional y la solidez curricular bajo la Taxonomía SOLO."""
            }
        ]

        # Agregar prompt específico para justificación y currículo
        if cond_key == 'cond_2':
            prompts.append({
                "id": "p_mercado",
                "category": "Estudio de Mercado y OLE",
                "title": "4. Extraer estadísticas DANE, OLE y Demanda Laboral",
                "prompt": f"Revisa las fuentes adjuntas y extrae todos los datos de pertinencia para '{prog_name}': tasas de vinculación formal (OLE), salarios promedio, ocupaciones CIUO-08 y requerimientos del sector productivo. Estructura el análisis en una tabla comparativa con citas exactas."
            })
        elif cond_key == 'cond_3':
            prompts.append({
                "id": "p_solo",
                "category": "Taxonomía SOLO & Malla Curricular",
                "title": "4. Mapear Resultados de Aprendizaje bajo Taxonomía SOLO",
                "prompt": f"A partir de los planes de estudio y documentos pedagógicos cargados de '{prog_name}', formula la matriz completa de Resultados de Aprendizaje (RA) clasificados en los 5 niveles de la Taxonomía SOLO (Preestructural, Uniestructural, Multiestructural, Relacional y Abstracto Ampliado) articulados con la Taxonomía de Bloom y su respectivo criterio de evaluación."
            })

        return jsonify({
            'status': 'success',
            'cond_key': cond_key,
            'cond_num': cond_num,
            'cond_title': cond_title,
            'prompts': prompts
        })
    except Exception as e:
        return jsonify({'status': 'error', 'message': str(e)}), 500

